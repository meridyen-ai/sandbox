"""
REST API for Sandbox

Provides HTTP endpoints for sandbox operations.
Alternative to gRPC for simpler integrations.
"""

from __future__ import annotations

import asyncio
import hashlib
import os
import secrets
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, date, time, timedelta, timezone
from decimal import Decimal
from enum import Enum
from ipaddress import IPv4Address, IPv6Address, IPv4Network, IPv6Network
from pathlib import Path
from typing import Any, AsyncIterator
from uuid import UUID

from fastapi import FastAPI, Cookie, File, Form, HTTPException, Depends, Header, Request, Response, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
import jwt

from sandbox.core.config import get_config
from sandbox.core.exceptions import (
    SandboxError,
    AuthenticationError,
    AuthorizationError,
    ValidationError,
)
from sandbox.core.logging import get_logger, bind_context, clear_context, setup_logging
from sandbox.execution.base import ExecutionContext
from sandbox.execution.sql_executor import SQLExecutor
from sandbox.execution.python_executor import PythonExecutor
from sandbox.visualization.generator import VisualizationGenerator, ChartType

logger = get_logger(__name__)


def _make_json_safe(value: Any) -> Any:
    """Convert any database value to a JSON-serializable type.

    Handles types from all supported databases (PostgreSQL, MySQL, MSSQL,
    Snowflake, BigQuery, etc.) so the API response is database-agnostic.
    """
    if value is None:
        return None
    # Primitives — fast path
    if isinstance(value, (bool, int, float, str)):
        return value
    # Date/time types (all databases)
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, time):
        return value.isoformat()
    if isinstance(value, timedelta):
        return value.total_seconds()
    # Numeric types
    if isinstance(value, Decimal):
        return float(value)
    # UUID (PostgreSQL, etc.)
    if isinstance(value, UUID):
        return str(value)
    # Binary data
    if isinstance(value, (bytes, bytearray, memoryview)):
        if isinstance(value, memoryview):
            value = bytes(value)
        return value.decode("utf-8", errors="replace")
    # Network types (PostgreSQL)
    if isinstance(value, (IPv4Address, IPv6Address, IPv4Network, IPv6Network)):
        return str(value)
    # Enum types
    if isinstance(value, Enum):
        return value.value
    # Path
    if isinstance(value, Path):
        return str(value)
    # Collections
    if isinstance(value, dict):
        return {str(k): _make_json_safe(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [_make_json_safe(v) for v in value]
    if isinstance(value, (set, frozenset)):
        return [_make_json_safe(v) for v in value]
    # Fallback — convert anything else to string
    try:
        return str(value)
    except Exception:
        return repr(value)


# =============================================================================
# Request/Response Models
# =============================================================================


class ExecutionContextModel(BaseModel):
    """Execution context from request."""
    request_id: str | None = Field(default_factory=lambda: str(uuid.uuid4()))
    workspace_id: str | None = None
    connection_id: str | None = None
    user_id: str | None = None
    max_rows: int | None = None
    timeout_seconds: int | None = None
    max_memory_mb: int | None = None
    max_output_size_kb: int | None = None


class SQLExecutionRequest(BaseModel):
    """SQL execution request."""
    context: ExecutionContextModel = Field(default_factory=ExecutionContextModel)
    query: str
    parameters: dict[str, Any] | None = None


class PythonExecutionRequest(BaseModel):
    """Python execution request."""
    context: ExecutionContextModel = Field(default_factory=ExecutionContextModel)
    code: str
    input_data: dict[str, Any] | None = None
    variables: dict[str, Any] | None = None


class VisualizationRequest(BaseModel):
    """Visualization request."""
    context: ExecutionContextModel = Field(default_factory=ExecutionContextModel)
    instruction: str | None = None
    data: list[dict[str, Any]]
    chart_type: str = "auto"
    title: str | None = None


class ConnectionConfig(BaseModel):
    """Database connection configuration."""
    id: str | None = None
    name: str
    db_type: str
    host: str
    port: int
    database: str
    username: str
    password: str
    schema_name: str | None = None
    ssl_enabled: bool = True

    @property
    def normalized_db_type(self) -> str:
        """Normalize common db_type aliases to canonical enum values."""
        aliases = {"postgres": "postgresql", "pg": "postgresql", "mssql_server": "mssql", "sqlserver": "mssql"}
        return aliases.get(self.db_type.lower(), self.db_type.lower())


class AIGenerateQueryRequest(BaseModel):
    """AI query generation request."""
    connection_id: str
    user_query: str


class GoogleSheetUploadRequest(BaseModel):
    """Request body for Google Sheets upload."""
    name: str
    spreadsheet_id: str
    credentials_json: str
    worksheet_name: str | None = None


class LoginRequest(BaseModel):
    """Login request with username and password."""
    username: str
    password: str


class LoginResponse(BaseModel):
    """Login response."""
    username: str
    message: str = "Login successful"


class SaveApiKeyRequest(BaseModel):
    """Request to save a Meridyen platform API key."""
    api_key: str = Field(..., min_length=4, max_length=200)


class HealthResponse(BaseModel):
    """Health check response."""
    status: str
    version: str
    uptime_seconds: float
    components: dict[str, dict[str, Any]] | None = None


class CapabilitiesResponse(BaseModel):
    """Capabilities response."""
    sandbox_id: str | None
    version: str
    supported_databases: list[str]
    supported_packages: list[str]
    resource_limits: dict[str, Any]
    supports_streaming: bool
    supports_visualization: bool
    has_local_llm: bool


# =============================================================================
# Dependencies
# =============================================================================


def _get_user_jwt_secret() -> str:
    """Get the JWT secret for user session tokens."""
    return os.environ.get("SANDBOX_AUTH_JWT_SECRET", "sandbox-jwt-secret-change-me-to-something-secure")


def _create_user_token(username: str) -> str:
    """Create a JWT token for an authenticated user."""
    secret = _get_user_jwt_secret()
    payload = {
        "sub": username,
        "type": "user_session",
        "iat": int(datetime.now(timezone.utc).timestamp()),
    }
    return jwt.encode(payload, secret, algorithm="HS256")


def _verify_user_token(token: str) -> dict[str, Any] | None:
    """Verify a user session JWT token. Returns claims or None."""
    secret = _get_user_jwt_secret()
    try:
        return jwt.decode(token, secret, algorithms=["HS256"])
    except jwt.InvalidTokenError:
        return None


# =============================================================================
# API Key Database Helpers
# =============================================================================

_api_key_engine = None


def _get_api_key_engine():
    """Get or create SQLAlchemy engine for the API keys database."""
    global _api_key_engine
    if _api_key_engine is None:
        from sqlalchemy import create_engine
        host = os.environ.get("SANDBOX_UPLOAD_DB_HOST", "sandbox-postgres")
        port = int(os.environ.get("SANDBOX_UPLOAD_DB_PORT", "5432"))
        db = os.environ.get("SANDBOX_UPLOAD_DB_NAME", "sandbox_uploads")
        user = os.environ.get("SANDBOX_UPLOAD_DB_USER", "sandbox")
        password = os.environ.get("SANDBOX_UPLOAD_DB_PASSWORD", "sandbox_password")
        url = f"postgresql+psycopg2://{user}:{password}@{host}:{port}/{db}"
        _api_key_engine = create_engine(url, pool_pre_ping=True, pool_size=3)
    return _api_key_engine


def _init_api_keys_table():
    """Create the api_keys table if it doesn't exist."""
    from sqlalchemy import text
    engine = _get_api_key_engine()
    with engine.connect() as conn:
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS api_keys (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                key_prefix TEXT NOT NULL,
                key_hash TEXT NOT NULL,
                workspace_id TEXT NOT NULL DEFAULT 'default',
                workspace_name TEXT NOT NULL DEFAULT 'Default Workspace',
                permissions JSONB NOT NULL DEFAULT '{"execute_sql": true, "execute_python": true, "generate_visualizations": true}',
                created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
            )
        """))
        conn.commit()
    logger.info("api_keys_table_initialized")


def _hash_api_key(raw_key: str) -> str:
    """Hash an API key with SHA-256."""
    return hashlib.sha256(raw_key.encode()).hexdigest()


def _generate_api_key() -> str:
    """Generate a new sb_ prefixed API key."""
    return "sb_" + secrets.token_urlsafe(32)


async def verify_sandbox_token(
    request: Request,
    authorization: str | None = Header(None, alias="Authorization"),
    x_api_key: str | None = Header(None, alias="X-API-Key"),
) -> dict[str, Any]:
    """
    Verify sandbox authentication token.

    Supports three authentication methods:
    1. User session cookie (sandbox_token) - For browser-based users
       - Cookie set by /api/v1/auth/login
       - Contains JWT with user session info

    2. Sandbox API Key (sb_*) - For programmatic access
       - Header: Authorization: Bearer sb_xxx OR X-API-Key: sb_xxx
       - Validates against configured auth provider

    3. JWT Token (legacy) - For platform-to-sandbox communication
       - Header: Authorization: Bearer <jwt_token>
       - Used for internal platform communication

    Returns:
        Dict with workspace_id, user_id, and other context
    """
    from sandbox.auth.sandbox_auth import get_auth_provider

    config = get_config()

    # 1. Try cookie-based user session first
    session_token = request.cookies.get("sandbox_token")
    if session_token:
        claims = _verify_user_token(session_token)
        if claims and claims.get("type") == "user_session":
            return {
                "auth_type": "user_session",
                "workspace_id": "default",
                "workspace_name": "Default Workspace",
                "user_id": claims.get("sub"),
                "permissions": {
                    "execute_sql": True,
                    "execute_python": True,
                    "generate_visualizations": True,
                },
            }

    # 2. Try X-API-Key header
    api_key = x_api_key

    # Fall back to Authorization header
    if not api_key and authorization:
        if authorization.startswith("Bearer "):
            api_key = authorization[7:]
        else:
            raise AuthenticationError("Invalid authorization format. Use 'Bearer <token>'")

    if not api_key:
        raise AuthenticationError("Authentication required. Provide credentials via login or X-API-Key header")

    # Check if it's a sandbox API key (sb_* prefix)
    if api_key.startswith("sb_"):
        provider = get_auth_provider()
        if not provider:
            raise AuthenticationError("Auth provider not initialized")

        auth_result = await provider.verify(api_key)
        if not auth_result:
            raise AuthenticationError("Invalid or inactive sandbox API key")

        # Return workspace context
        return {
            "auth_type": "sandbox_api_key",
            "workspace_id": str(auth_result.workspace_id) if auth_result.workspace_id else None,
            "workspace_name": auth_result.workspace_name,
            "user_id": str(auth_result.user_id) if auth_result.user_id else None,
            "api_key_name": auth_result.api_key_name,
            "permissions": auth_result.permissions or {
                "execute_sql": True,
                "execute_python": True,
                "generate_visualizations": True,
            },
        }

    # 3. Otherwise, try to decode as JWT (legacy method for platform communication)
    try:
        secret = config.platform.registration_token
        if secret:
            payload = jwt.decode(
                api_key,
                secret.get_secret_value(),
                algorithms=["HS256"],
                audience="sandbox-executor",
            )
            payload["auth_type"] = "jwt"
            return payload
        else:
            # Development mode - accept any token
            logger.warning("Development mode: accepting token without verification")
            return {
                "auth_type": "dev",
                "workspace_id": "dev",
                "permissions": {}
            }

    except jwt.ExpiredSignatureError:
        raise AuthenticationError("Token expired")
    except jwt.InvalidTokenError as e:
        raise AuthenticationError(f"Invalid token: {e}")


# =============================================================================
# Application Factory
# =============================================================================


def create_rest_app() -> FastAPI:
    """Create and configure the FastAPI application."""
    config = get_config()

    # Lifespan manager
    @asynccontextmanager
    async def lifespan(app: FastAPI) -> AsyncIterator[None]:
        # Startup
        setup_logging()
        logger.info("rest_api_starting", port=config.server.rest_port)

        # Initialize auth provider if API key auth is enabled
        if config.authentication.enable_api_key_auth:
            from sandbox.auth.sandbox_auth import initialize_auth_provider
            try:
                initialize_auth_provider(config)
                logger.info(f"Auth provider initialized: {config.authentication.provider}")
            except Exception as e:
                logger.error(f"Failed to initialize auth provider: {e}")
                if config.environment == "production":
                    raise

        # Initialize API keys table
        try:
            _init_api_keys_table()
        except Exception as e:
            logger.error(f"Failed to initialize api_keys table: {e}")

        # Initialize executors
        app.state.sql_executor = SQLExecutor()
        app.state.python_executor = PythonExecutor()
        app.state.viz_generator = VisualizationGenerator()
        app.state.start_time = datetime.now(timezone.utc)

        yield

        # Shutdown
        logger.info("rest_api_stopping")
        await app.state.sql_executor.close()

        # Close auth provider
        if config.authentication.enable_api_key_auth:
            from sandbox.auth.sandbox_auth import get_auth_provider
            provider = get_auth_provider()
            if provider:
                await provider.close()

    app = FastAPI(
        title="Meridyen Sandbox API",
        description="Secure execution sandbox for SQL and Python code",
        version="1.0.0",
        lifespan=lifespan,
        docs_url="/docs" if config.debug else None,
        redoc_url="/redoc" if config.debug else None,
    )

    # CORS middleware — allow all origins for sandbox (API key or cookie auth handles security)
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    # Exception handlers
    @app.exception_handler(SandboxError)
    async def sandbox_error_handler(request, exc: SandboxError) -> JSONResponse:
        status_code = status.HTTP_400_BAD_REQUEST
        if isinstance(exc, AuthenticationError):
            status_code = status.HTTP_401_UNAUTHORIZED
        elif isinstance(exc, AuthorizationError):
            status_code = status.HTTP_403_FORBIDDEN

        return JSONResponse(
            status_code=status_code,
            content=exc.to_dict(),
        )

    # Routes
    register_routes(app)

    # Serve built frontend static files (production mode)
    # The frontend build output goes to frontend/dist
    frontend_dist = Path(__file__).resolve().parent.parent.parent.parent / "frontend" / "dist"
    if frontend_dist.exists() and (frontend_dist / "index.html").exists():
        # Mount static assets (js, css, images, etc.)
        app.mount("/assets", StaticFiles(directory=str(frontend_dist / "assets")), name="static-assets")
        # Mount icons if they exist
        icons_dir = frontend_dist / "icons"
        if icons_dir.exists():
            app.mount("/icons", StaticFiles(directory=str(icons_dir)), name="static-icons")

        # Catch-all: serve index.html for client-side routing
        @app.get("/{path:path}", include_in_schema=False)
        async def serve_spa(path: str) -> FileResponse:
            # Resolve and verify the path stays within frontend_dist
            file_path = (frontend_dist / path).resolve()
            if file_path.is_file() and str(file_path).startswith(str(frontend_dist.resolve())):
                return FileResponse(str(file_path))
            # Otherwise serve index.html for client-side routing
            return FileResponse(str(frontend_dist / "index.html"))

        logger.info("serving_frontend", dist_path=str(frontend_dist))

    return app


def register_routes(app: FastAPI) -> None:
    """Register API routes."""

    # ==========================================================================
    # Health & Capabilities
    # ==========================================================================

    @app.get("/health", response_model=HealthResponse, tags=["Health"])
    async def health_check(include_details: bool = False) -> HealthResponse:
        """Health check endpoint."""
        uptime = (datetime.now(timezone.utc) - app.state.start_time).total_seconds()

        response = HealthResponse(
            status="healthy",
            version="1.0.0",
            uptime_seconds=uptime,
        )

        if include_details:
            response.components = {
                "sql_executor": {"status": "healthy"},
                "python_executor": {"status": "healthy"},
                "visualization": {"status": "healthy"},
            }

        return response

    @app.get("/capabilities", response_model=CapabilitiesResponse, tags=["Health"])
    async def get_capabilities() -> CapabilitiesResponse:
        """Get sandbox capabilities."""
        from sandbox.connectors.factory import get_available_connectors

        config = get_config()

        return CapabilitiesResponse(
            sandbox_id=config.platform.sandbox_id,
            version="1.0.0",
            supported_databases=get_available_connectors(),
            supported_packages=list(config.security.allowed_python_imports),
            resource_limits={
                "max_memory_mb": config.resource_limits.max_memory_mb,
                "max_cpu_seconds": config.resource_limits.max_cpu_seconds,
                "max_output_size_kb": config.resource_limits.max_output_size_kb,
                "max_rows": config.resource_limits.max_rows,
                "query_timeout_seconds": config.resource_limits.query_timeout_seconds,
                "python_timeout_seconds": config.resource_limits.python_timeout_seconds,
            },
            supports_streaming=True,
            supports_visualization=True,
            has_local_llm=config.local_llm.enabled,
        )

    # ==========================================================================
    # User Authentication (username/password from .env)
    # ==========================================================================

    @app.post("/api/v1/auth/login", tags=["Auth"])
    async def auth_login(payload: LoginRequest, response: Response) -> JSONResponse:
        """Authenticate with username and password defined in .env."""
        expected_username = os.environ.get("SANDBOX_AUTH_USERNAME", "admin")
        expected_password = os.environ.get("SANDBOX_AUTH_PASSWORD", "admin123")

        if (
            payload.username.lower() == expected_username.lower()
            and payload.password == expected_password
        ):
            token = _create_user_token(payload.username)
            response = JSONResponse(
                content={
                    "username": payload.username,
                    "message": "Login successful",
                }
            )
            response.set_cookie(
                key="sandbox_token",
                value=token,
                httponly=True,
                secure=os.environ.get("SANDBOX_ENVIRONMENT", "") in ("production", "preprod"),
                samesite="lax",
                max_age=86400 * 7,  # 7 days
                path="/",
            )
            return response

        raise HTTPException(status_code=401, detail="Invalid username or password")

    @app.get("/api/v1/auth/me", tags=["Auth"])
    async def auth_me(request: Request) -> JSONResponse:
        """Check current user session from cookie."""
        session_token = request.cookies.get("sandbox_token")
        if not session_token:
            raise HTTPException(status_code=401, detail="Not authenticated")

        claims = _verify_user_token(session_token)
        if not claims or claims.get("type") != "user_session":
            raise HTTPException(status_code=401, detail="Invalid session")

        return JSONResponse(content={
            "username": claims.get("sub"),
            "authenticated": True,
        })

    @app.post("/api/v1/auth/logout", tags=["Auth"])
    async def auth_logout(response: Response) -> JSONResponse:
        """Clear user session cookie."""
        response = JSONResponse(content={"message": "Logged out"})
        response.delete_cookie(
            key="sandbox_token",
            httponly=True,
            secure=os.environ.get("SANDBOX_ENVIRONMENT", "") in ("production", "preprod"),
            samesite="lax",
            path="/",
        )
        return response

    # ==========================================================================
    # API Key Management (single key per sandbox, provided by MVP platform)
    # ==========================================================================

    @app.get("/api/v1/api-key", tags=["API Keys"])
    async def get_api_key(
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Get the current sandbox API key info (prefix only, never the full key)."""
        from sqlalchemy import text

        engine = _get_api_key_engine()
        with engine.connect() as conn:
            result = conn.execute(text(
                "SELECT id, key_prefix, created_at FROM api_keys ORDER BY created_at DESC LIMIT 1"
            ))
            row = result.fetchone()

        if not row:
            return JSONResponse(content={"configured": False})

        return JSONResponse(content={
            "configured": True,
            "key_prefix": row[1],
            "created_at": row[2].isoformat() if row[2] else None,
        })

    @app.put("/api/v1/api-key", tags=["API Keys"])
    async def save_api_key(
        payload: SaveApiKeyRequest,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Save the Meridyen platform API key. Replaces any existing key."""
        from sqlalchemy import text
        import json

        api_key = payload.api_key.strip()
        if not api_key.startswith("sb_"):
            raise HTTPException(status_code=400, detail="API key must start with 'sb_'")

        key_id = str(uuid.uuid4())
        key_prefix = api_key[:13]  # "sb_" + first 10 chars
        key_hash = _hash_api_key(api_key)
        permissions = {"execute_sql": True, "execute_python": True, "generate_visualizations": True}

        engine = _get_api_key_engine()
        with engine.connect() as conn:
            # Replace any existing key — sandbox has only one
            conn.execute(text("DELETE FROM api_keys"))
            conn.execute(
                text("""
                    INSERT INTO api_keys (id, name, key_prefix, key_hash, workspace_id, workspace_name, permissions)
                    VALUES (:id, :name, :key_prefix, :key_hash, :workspace_id, :workspace_name, CAST(:permissions AS jsonb))
                """),
                {
                    "id": key_id,
                    "name": "meridyen-platform-key",
                    "key_prefix": key_prefix,
                    "key_hash": key_hash,
                    "workspace_id": "default",
                    "workspace_name": "Default Workspace",
                    "permissions": json.dumps(permissions),
                },
            )
            conn.commit()

        logger.info("api_key_saved", key_prefix=key_prefix)

        return JSONResponse(content={
            "configured": True,
            "key_prefix": key_prefix,
        })

    @app.delete("/api/v1/api-key", tags=["API Keys"])
    async def remove_api_key(
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Remove the configured API key."""
        from sqlalchemy import text

        engine = _get_api_key_engine()
        with engine.connect() as conn:
            conn.execute(text("DELETE FROM api_keys"))
            conn.commit()

        logger.info("api_key_removed")
        return JSONResponse(content={"message": "API key removed"})

    # ==========================================================================
    # Execution Endpoints
    # ==========================================================================

    @app.post("/api/v1/execute/sql", tags=["Execution"])
    async def execute_sql(
        request: SQLExecutionRequest,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Execute SQL query."""
        request_id = request.context.request_id or str(uuid.uuid4())
        bind_context(request_id=request_id)

        try:
            # Build execution context
            exec_context = ExecutionContext(
                request_id=request_id,
                workspace_id=request.context.workspace_id,
                connection_id=request.context.connection_id,
                user_id=request.context.user_id,
                max_rows=request.context.max_rows,
                timeout_seconds=request.context.timeout_seconds,
            )

            # Validate
            errors = await app.state.sql_executor.validate(exec_context, query=request.query)
            if errors:
                raise ValidationError("; ".join(errors))

            # Execute
            result = await app.state.sql_executor.execute(
                exec_context,
                query=request.query,
                parameters=request.parameters,
            )

            return JSONResponse(
                content={
                    "request_id": request_id,
                    "status": "success" if result.is_success() else "error",
                    "data": {
                        "columns": [
                            {"name": c.name, "type": c.data_type, "masked": c.is_masked}
                            for c in result.columns
                        ],
                        "rows": [_make_json_safe(row) for row in result.rows],
                        "row_count": result.row_count,
                        "total_rows_available": result.total_rows_available,
                    },
                    "metrics": result.metrics.to_dict(),
                }
            )

        except SandboxError:
            raise
        except Exception as e:
            logger.error("execute_sql_error", request_id=request_id, error=str(e))
            raise HTTPException(status_code=500, detail=str(e))
        finally:
            clear_context()

    @app.post("/api/v1/execute/python", tags=["Execution"])
    async def execute_python(
        request: PythonExecutionRequest,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Execute Python code."""
        request_id = request.context.request_id or str(uuid.uuid4())
        bind_context(request_id=request_id)

        try:
            exec_context = ExecutionContext(
                request_id=request_id,
                workspace_id=request.context.workspace_id,
                user_id=request.context.user_id,
                timeout_seconds=request.context.timeout_seconds,
                max_memory_mb=request.context.max_memory_mb,
                max_output_size_kb=request.context.max_output_size_kb,
            )

            # Validate
            errors = await app.state.python_executor.validate(exec_context, code=request.code)
            if errors:
                raise ValidationError("; ".join(errors))

            # Build input data
            input_data = {}
            if request.input_data:
                input_data["data"] = request.input_data.get("data", [])
            if request.variables:
                input_data["variables"] = request.variables

            # Execute
            result = await app.state.python_executor.execute(
                exec_context,
                code=request.code,
                input_data=input_data,
            )

            return JSONResponse(
                content={
                    "request_id": request_id,
                    "status": "success" if result.is_success() else "error",
                    "data": {
                        "stdout": result.stdout,
                        "stderr": result.stderr,
                        "result": result.result_data,
                        "variables": result.variables,
                    },
                    "metrics": result.metrics.to_dict(),
                    "error": result.error_message if not result.is_success() else None,
                }
            )

        except SandboxError:
            raise
        except Exception as e:
            logger.error("execute_python_error", request_id=request_id, error=str(e))
            raise HTTPException(status_code=500, detail=str(e))
        finally:
            clear_context()

    @app.post("/api/v1/visualize", tags=["Visualization"])
    async def create_visualization(
        request: VisualizationRequest,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Create visualization from data."""
        request_id = request.context.request_id or str(uuid.uuid4())
        bind_context(request_id=request_id)

        try:
            exec_context = ExecutionContext(
                request_id=request_id,
                workspace_id=request.context.workspace_id,
                max_output_size_kb=request.context.max_output_size_kb,
            )

            # Map chart type
            chart_type_map = {
                "auto": ChartType.AUTO,
                "line": ChartType.LINE,
                "bar": ChartType.BAR,
                "pie": ChartType.PIE,
                "scatter": ChartType.SCATTER,
                "heatmap": ChartType.HEATMAP,
                "table": ChartType.TABLE,
            }
            chart_type = chart_type_map.get(request.chart_type.lower(), ChartType.AUTO)

            # Generate visualization
            result = await app.state.viz_generator.generate(
                exec_context,
                data=request.data,
                instruction=request.instruction,
                chart_type=chart_type,
                title=request.title,
            )

            return JSONResponse(
                content={
                    "request_id": request_id,
                    "status": "success" if result.status.value == "success" else "error",
                    "data": {
                        "plotly_spec": result.plotly_spec,
                        "insight": result.insight,
                        "explanation": result.explanation,
                        "chart_type": result.chart_type.value,
                        "data_points": result.data_points,
                    },
                    "metrics": result.metrics.to_dict(),
                    "error": result.error_message,
                }
            )

        except SandboxError:
            raise
        except Exception as e:
            logger.error("create_visualization_error", request_id=request_id, error=str(e))
            raise HTTPException(status_code=500, detail=str(e))
        finally:
            clear_context()

    # ==========================================================================
    # Handlers (Database Types)
    # ==========================================================================

    @app.get("/api/v1/handlers", tags=["Handlers"])
    async def list_handlers() -> JSONResponse:
        """
        List all available database handlers.

        Returns handler metadata including connection arguments for dynamic form generation.
        """
        from sandbox.services.db_handler_service import DBHandlerService

        handlers = DBHandlerService.get_available_handlers()

        return JSONResponse(content={
            "handlers": [handler.to_dict() for handler in handlers]
        })

    # ==========================================================================
    # Connection Management
    # ==========================================================================

    @app.get("/api/v1/connections", tags=["Connections"])
    async def list_connections(
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """List configured database connections."""
        from sandbox.core.connection_store import list_connections as db_list_connections

        rows = db_list_connections()
        connections = []
        for row in rows:
            connections.append({
                "id": row["id"],
                "name": row["name"],
                "db_type": row["db_type"],
                "host": row["host"],
                "port": row["port"],
                "database": row["database"],
                "schema": row.get("schema_name"),
                "is_default": False,
                "created_at": row.get("created_at"),
                "updated_at": row.get("updated_at"),
            })

        return JSONResponse(content={"connections": connections})

    @app.get("/api/v1/connections/{connection_id}", tags=["Connections"])
    async def get_connection(
        connection_id: str,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Return one connection including credentials.

        Used by MVP schema sync when the workspace mirror row does not store secrets.
        Requires the same API key / auth as other sandbox routes.
        """
        from sandbox.core.connection_store import get_connection as db_get_connection

        row = db_get_connection(connection_id)
        if not row:
            raise HTTPException(status_code=404, detail="Connection not found")

        return JSONResponse(
            content={
                "id": row["id"],
                "name": row["name"],
                "db_type": row["db_type"],
                "host": row["host"],
                "port": row["port"],
                "database": row["database"],
                "username": row["username"],
                "password": row["password"],
                "schema_name": row.get("schema_name"),
                "ssl_enabled": row.get("ssl_enabled", False),
            }
        )

    @app.post("/api/v1/connections", tags=["Connections"])
    async def create_connection(
        connection: ConnectionConfig,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Create a new database connection."""
        from sandbox.core.connection_store import create_connection as db_create_connection
        from sandbox.core.config import DatabaseConnectionConfig, DatabaseType, get_config
        from pydantic import SecretStr
        import uuid

        conn_id = connection.id or str(uuid.uuid4())

        row = db_create_connection({
            "id": conn_id,
            "name": connection.name,
            "db_type": connection.normalized_db_type,
            "host": connection.host,
            "port": connection.port,
            "database": connection.database,
            "username": connection.username,
            "password": connection.password,
            "schema_name": connection.schema_name,
            "ssl_enabled": connection.ssl_enabled,
        })

        # Also add to in-memory config so queries work immediately
        config = get_config()
        new_conn = DatabaseConnectionConfig(
            id=conn_id,
            name=connection.name,
            db_type=DatabaseType(connection.normalized_db_type),
            host=connection.host,
            port=connection.port,
            database=connection.database,
            username=connection.username,
            password=SecretStr(connection.password),
            schema_name=connection.schema_name,
            ssl_enabled=connection.ssl_enabled,
            created_at=row["created_at"] if row else None,
            updated_at=row["updated_at"] if row else None,
        )
        # Remove existing with same id if any
        config.database_connections = [c for c in config.database_connections if c.id != conn_id]
        config.database_connections.append(new_conn)

        return JSONResponse(
            status_code=201,
            content={
                "id": conn_id,
                "name": connection.name,
                "message": "Connection created successfully"
            }
        )

    @app.put("/api/v1/connections/{connection_id}", tags=["Connections"])
    async def update_connection(
        connection_id: str,
        connection: ConnectionConfig,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Update an existing database connection."""
        from sandbox.core.connection_store import update_connection as db_update_connection
        from sandbox.core.config import DatabaseConnectionConfig, DatabaseType, get_config
        from pydantic import SecretStr

        row = db_update_connection(connection_id, {
            "name": connection.name,
            "db_type": connection.normalized_db_type,
            "host": connection.host,
            "port": connection.port,
            "database": connection.database,
            "username": connection.username,
            "password": connection.password,
            "schema_name": connection.schema_name,
            "ssl_enabled": connection.ssl_enabled,
        })

        if not row:
            raise HTTPException(status_code=404, detail="Connection not found")

        # Update in-memory config
        config = get_config()
        for idx, conn in enumerate(config.database_connections):
            if conn.id == connection_id:
                config.database_connections[idx] = DatabaseConnectionConfig(
                    id=connection_id,
                    name=connection.name,
                    db_type=DatabaseType(connection.normalized_db_type),
                    host=connection.host,
                    port=connection.port,
                    database=connection.database,
                    username=connection.username,
                    password=SecretStr(connection.password),
                    schema_name=connection.schema_name,
                    ssl_enabled=connection.ssl_enabled,
                    created_at=row.get("created_at"),
                    updated_at=row.get("updated_at"),
                )
                break

        return JSONResponse(content={
            "id": connection_id,
            "message": "Connection updated successfully"
        })

    @app.delete("/api/v1/connections/{connection_id}", tags=["Connections"])
    async def delete_connection(
        connection_id: str,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Delete a database connection."""
        from sandbox.core.connection_store import delete_connection as db_delete_connection

        deleted = db_delete_connection(connection_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="Connection not found")

        # Remove from in-memory config
        config = get_config()
        config.database_connections = [c for c in config.database_connections if c.id != connection_id]

        return JSONResponse(content={
            "message": "Connection deleted successfully"
        })

    @app.get("/api/v1/connections/{connection_id}/selected-tables", tags=["Connections"])
    async def get_selected_tables(
        connection_id: str,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Get selected tables/columns for a connection."""
        from sandbox.core.connection_store import get_selected_tables as db_get_selected_tables

        tables = db_get_selected_tables(connection_id)
        if tables is None:
            raise HTTPException(status_code=404, detail="Connection not found")

        return JSONResponse(content={
            "connection_id": connection_id,
            "selected_tables": tables,
        })

    @app.put("/api/v1/connections/{connection_id}/selected-tables", tags=["Connections"])
    async def save_selected_tables(
        connection_id: str,
        payload: dict,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Save selected tables/columns for a connection."""
        from sandbox.core.connection_store import save_selected_tables as db_save_selected_tables

        selected = payload.get("selected_tables", {})
        updated = db_save_selected_tables(connection_id, selected)
        if not updated:
            raise HTTPException(status_code=404, detail="Connection not found")

        # Update in-memory config
        config = get_config()
        for conn in config.database_connections:
            if conn.id == connection_id:
                conn.selected_tables = selected
                break

        return JSONResponse(content={
            "connection_id": connection_id,
            "message": "Selected tables saved successfully",
            "selected_tables": selected,
        })

    @app.post("/api/v1/connections/test", tags=["Connections"])
    async def test_connection(
        connection: ConnectionConfig,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Test a database connection."""
        from sandbox.connectors.factory import get_connector
        from sandbox.core.config import DatabaseConnectionConfig, DatabaseType
        from pydantic import SecretStr

        try:
            # Build config
            conn_config = DatabaseConnectionConfig(
                id=connection.id or str(uuid.uuid4()),
                name=connection.name,
                db_type=DatabaseType(connection.normalized_db_type),
                host=connection.host,
                port=connection.port,
                database=connection.database,
                username=connection.username,
                password=SecretStr(connection.password),
                schema_name=connection.schema_name,
                ssl_enabled=connection.ssl_enabled,
            )

            # Get connector and test
            connector = get_connector(conn_config.db_type, conn_config)
            conn = await connector.connect()
            is_valid = await connector.test_connection(conn)
            await connector.close_connection(conn)

            return JSONResponse(
                content={
                    "success": is_valid,
                    "message": "Connection successful" if is_valid else "Connection test failed",
                }
            )

        except Exception as e:
            return JSONResponse(
                content={
                    "success": False,
                    "message": str(e),
                }
            )

    # ==========================================================================
    # File Upload (CSV, XLSX, XLS)
    # ==========================================================================

    FILE_UPLOAD_DIR = Path("/app/data/uploads")
    ALLOWED_FILE_EXTENSIONS = {".csv", ".xlsx", ".xls"}

    @app.post("/api/v1/upload-file", tags=["File Upload"])
    async def upload_file(
        file: UploadFile = File(...),
        name: str = Form(...),
        delimiter: str = Form(","),
        has_header: str = Form("true"),
        selected_sheets: str = Form(""),
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Upload a CSV, XLSX, or XLS file and load it into the sandbox PostgreSQL.
        Each CSV becomes a table; each Excel sheet becomes a separate table.
        Data is queryable via standard SQL after upload.
        """
        import os
        import pandas as pd
        from sandbox.core.config import (
            DatabaseConnectionConfig, DatabaseType, get_config,
        )
        from sandbox.core.connection_store import create_connection as db_create_connection
        from sandbox.services.file_loader import (
            create_upload_database, sanitize_table_name,
            load_csv_to_postgres, load_excel_sheet_to_postgres,
        )
        from pydantic import SecretStr

        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in ALLOWED_FILE_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type '{file_ext}'. Allowed: CSV, XLSX, XLS",
            )

        is_excel = file_ext in (".xlsx", ".xls")

        FILE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        unique_id = str(uuid.uuid4())[:8]
        safe_name = "".join(c for c in name if c.isalnum() or c in "._- ").strip() or "uploaded_data"
        filename = f"{safe_name}_{unique_id}{file_ext}"
        file_path = str(FILE_UPLOAD_DIR / filename)

        # Stream file to disk in chunks (handles large files without loading into memory)
        try:
            with open(file_path, "wb") as f:
                while chunk := await file.read(1024 * 1024):  # 1MB chunks
                    f.write(chunk)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")

        has_header_bool = has_header.lower() == "true"
        config = get_config()
        now = datetime.now(timezone.utc).isoformat()

        # Create a dedicated database for this upload
        sql_engine, db_config = create_upload_database(name)

        try:

            if is_excel:
                excel_engine = "openpyxl" if file_ext == ".xlsx" else "xlrd"
                excel_file = pd.ExcelFile(file_path, engine=excel_engine)
                all_sheets = excel_file.sheet_names

                if not all_sheets:
                    os.remove(file_path)
                    raise HTTPException(status_code=400, detail="Excel file contains no sheets")

                if selected_sheets.strip():
                    sheets_to_import = [s.strip() for s in selected_sheets.split(",") if s.strip()]
                    invalid = [s for s in sheets_to_import if s not in all_sheets]
                    if invalid:
                        os.remove(file_path)
                        raise HTTPException(
                            status_code=400,
                            detail=f"Sheet(s) not found: {', '.join(invalid)}. Available: {', '.join(all_sheets)}",
                        )
                else:
                    sheets_to_import = all_sheets

                base_table = sanitize_table_name(name)
                connections = []
                total_rows = 0
                sheet_stats = []

                for sheet_name in sheets_to_import:
                    # Table name: base_sheetname for multi-sheet, just base for single
                    if len(sheets_to_import) > 1:
                        table_name = f"{base_table}_{sanitize_table_name(sheet_name)}"
                    else:
                        table_name = base_table

                    result = load_excel_sheet_to_postgres(
                        file_path, sheet_name, table_name,
                        has_header=has_header_bool, engine=sql_engine,
                    )

                    if result["row_count"] == 0:
                        continue

                    total_rows += result["row_count"]
                    sheet_stats.append({
                        "sheet_name": sheet_name,
                        "table_name": result["table_name"],
                        "row_count": result["row_count"],
                        "column_count": result["column_count"],
                    })

                if not sheet_stats:
                    os.remove(file_path)
                    raise HTTPException(status_code=400, detail="All sheets are empty")

                # Create ONE connection for the upload database
                conn_id = str(uuid.uuid4())
                db_create_connection({
                    "id": conn_id,
                    "name": name,
                    "db_type": "postgresql",
                    "host": db_config["host"],
                    "port": db_config["port"],
                    "database": db_config["database"],
                    "username": db_config["username"],
                    "password": db_config["password"],
                    "ssl_enabled": False,
                })
                new_conn = DatabaseConnectionConfig(
                    id=conn_id,
                    name=name,
                    db_type=DatabaseType.POSTGRESQL,
                    host=db_config["host"],
                    port=db_config["port"],
                    database=db_config["database"],
                    username=db_config["username"],
                    password=SecretStr(db_config["password"]),
                    ssl_enabled=False,
                    created_at=now,
                    updated_at=now,
                )
                config.database_connections.append(new_conn)

                # Build backward-compatible connections list
                connections = [
                    {
                        "connection_id": conn_id,
                        "name": f"{name} - {s['sheet_name']}" if len(sheet_stats) > 1 else name,
                        "sheet_name": s["sheet_name"],
                        "table_name": s["table_name"],
                        "row_count": s["row_count"],
                    }
                    for s in sheet_stats
                ]

                # Clean up source file after successful load
                if os.path.exists(file_path):
                    os.remove(file_path)

                return JSONResponse(
                    status_code=201,
                    content={
                        "success": True,
                        "message": f"Excel file loaded into PostgreSQL: {len(sheet_stats)} sheet(s), {total_rows} total rows",
                        "connection_id": conn_id,
                        "name": name,
                        "db_type": "csv",
                        "source_type": "excel",
                        "host": db_config["host"],
                        "port": db_config["port"],
                        "database": db_config["database"],
                        "upload_tables": [s["table_name"] for s in sheet_stats],
                        "connections": connections,
                        "row_count": total_rows,
                        "sheets": all_sheets,
                        "sheet_stats": sheet_stats,
                    },
                )
            else:
                # CSV — load into PostgreSQL
                table_name = sanitize_table_name(name)
                result = load_csv_to_postgres(
                    file_path, table_name,
                    delimiter=delimiter, has_header=has_header_bool,
                    engine=sql_engine,
                )

                if result["row_count"] == 0:
                    os.remove(file_path)
                    raise HTTPException(status_code=400, detail="CSV file is empty")

                conn_id = str(uuid.uuid4())
                db_create_connection({
                    "id": conn_id,
                    "name": name,
                    "db_type": "postgresql",
                    "host": db_config["host"],
                    "port": db_config["port"],
                    "database": db_config["database"],
                    "username": db_config["username"],
                    "password": db_config["password"],
                    "ssl_enabled": False,
                })
                new_conn = DatabaseConnectionConfig(
                    id=conn_id,
                    name=name,
                    db_type=DatabaseType.POSTGRESQL,
                    host=db_config["host"],
                    port=db_config["port"],
                    database=db_config["database"],
                    username=db_config["username"],
                    password=SecretStr(db_config["password"]),
                    ssl_enabled=False,
                    created_at=now,
                    updated_at=now,
                )
                config.database_connections.append(new_conn)

                # Clean up source file after successful load
                if os.path.exists(file_path):
                    os.remove(file_path)

                return JSONResponse(
                    status_code=201,
                    content={
                        "success": True,
                        "message": f"CSV loaded into PostgreSQL: {result['row_count']} rows, {result['column_count']} columns",
                        "connection_id": conn_id,
                        "name": name,
                        "db_type": "csv",
                        "source_type": "csv",
                        "host": db_config["host"],
                        "port": db_config["port"],
                        "database": db_config["database"],
                        "upload_tables": [result["table_name"]],
                        "table_name": result["table_name"],
                        "row_count": result["row_count"],
                        "column_count": result["column_count"],
                    },
                )

        except HTTPException:
            raise
        except Exception as e:
            import os as _os
            if _os.path.exists(file_path):
                _os.remove(file_path)
            raise HTTPException(status_code=400, detail=f"Failed to load file into database: {str(e)}")

    @app.post("/api/v1/upload-file/sheets", tags=["File Upload"])
    async def get_file_sheets(
        file: UploadFile = File(...),
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Preview sheet names from an uploaded Excel file without creating a connection.
        Used by the frontend to let users select which sheets to import.
        """
        import os
        import pandas as pd

        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in (".xlsx", ".xls"):
            raise HTTPException(status_code=400, detail="Sheet detection is only for Excel files")

        FILE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        temp_path = str(FILE_UPLOAD_DIR / f"_temp_{uuid.uuid4().hex[:8]}{file_ext}")

        try:
            content = await file.read()
            with open(temp_path, "wb") as f:
                f.write(content)

            engine = "openpyxl" if file_ext == ".xlsx" else "xlrd"
            excel_file = pd.ExcelFile(temp_path, engine=engine)
            sheets = []
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(temp_path, sheet_name=sheet_name, nrows=5, engine=engine)
                sheets.append({
                    "name": sheet_name,
                    "columns": [str(c) for c in df.columns],
                    "preview_rows": len(df),
                })

            return JSONResponse(content={"sheets": sheets})
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read Excel file: {str(e)}")
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    # ==========================================================================
    # Google Sheets Upload (loads sheet data into sandbox PostgreSQL)
    # ==========================================================================

    @app.post("/api/v1/upload-gsheet", tags=["File Upload"])
    async def upload_google_sheet(
        body: GoogleSheetUploadRequest,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Fetch data from a Google Sheet and load it into sandbox PostgreSQL.
        Works like CSV upload — data becomes a table queryable via standard SQL.
        """
        import json as _json
        import os
        import pandas as pd
        from sandbox.core.config import (
            DatabaseConnectionConfig, DatabaseType, get_config,
        )
        from sandbox.core.connection_store import create_connection as db_create_connection
        from sandbox.services.file_loader import (
            create_upload_database, sanitize_table_name,
            load_dataframe_to_postgres, load_csv_to_postgres,
            load_excel_sheet_to_postgres,
        )
        from pydantic import SecretStr

        # Parse credentials
        try:
            if isinstance(body.credentials_json, str):
                credentials_info = _json.loads(body.credentials_json)
            else:
                credentials_info = body.credentials_json
        except _json.JSONDecodeError as e:
            raise HTTPException(status_code=400, detail=f"Invalid credentials JSON: {e}")

        # Connect to Google Sheets
        try:
            import gspread
            from google.oauth2.service_account import Credentials
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="Google Sheets support not installed (gspread, google-auth)",
            )

        try:
            scopes = [
                "https://www.googleapis.com/auth/spreadsheets.readonly",
                "https://www.googleapis.com/auth/drive.readonly",
            ]
            credentials = Credentials.from_service_account_info(credentials_info, scopes=scopes)
            client = gspread.authorize(credentials)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to authenticate with Google: {e}")

        # Try opening as native Google Sheet first; fall back to Drive download
        # for uploaded files (Excel/CSV opened in Google Sheets)
        is_native_sheet = True
        try:
            spreadsheet = client.open_by_key(body.spreadsheet_id)
            # Probe: listing worksheets fails on non-native files
            _ = spreadsheet.worksheets()
        except Exception:
            is_native_sheet = False

        config = get_config()
        now = datetime.now(timezone.utc).isoformat()

        # Create a dedicated database for this upload
        sql_engine, db_config = create_upload_database(body.name)

        try:
            base_table = sanitize_table_name(body.name)
            sheet_stats = []
            total_rows = 0

            if not is_native_sheet:
                # Non-native file in Drive (uploaded Excel/CSV) — download via
                # gspread's authorized session and reuse existing file loaders.
                import requests as _requests

                session = client.http_client.session
                download_url = f"https://www.googleapis.com/drive/v3/files/{body.spreadsheet_id}?alt=media"
                resp = session.get(download_url)
                if resp.status_code == 403:
                    sa_email = credentials_info.get("client_email", "the service account")
                    # Parse error details from Google API response
                    try:
                        error_info = resp.json()
                        error_msg = error_info.get("error", {}).get("message", "")
                    except Exception:
                        error_msg = ""
                    detail = f"Access denied. Share the file in Google Drive with: {sa_email}"
                    if error_msg:
                        detail += f" (Google: {error_msg})"
                    raise HTTPException(status_code=400, detail=detail)
                if resp.status_code != 200:
                    # Include response body for debugging
                    try:
                        error_body = resp.text[:500]
                    except Exception:
                        error_body = ""
                    raise HTTPException(
                        status_code=400,
                        detail=f"Failed to download file from Google Drive (HTTP {resp.status_code}): {error_body}",
                    )

                # Get file name to determine format
                meta_url = f"https://www.googleapis.com/drive/v3/files/{body.spreadsheet_id}?fields=name,mimeType"
                meta_resp = session.get(meta_url)
                file_name = "download"
                mime = ""
                if meta_resp.status_code == 200:
                    meta = meta_resp.json()
                    file_name = meta.get("name", "download")
                    mime = meta.get("mimeType", "")

                # Save to temp file and use existing file loaders
                FILE_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
                # Determine extension
                if file_name.endswith(".xlsx"):
                    ext = ".xlsx"
                elif file_name.endswith(".xls"):
                    ext = ".xls"
                else:
                    ext = ".csv"

                temp_path = str(FILE_UPLOAD_DIR / f"_gsheet_{uuid.uuid4().hex[:8]}{ext}")
                try:
                    with open(temp_path, "wb") as f:
                        f.write(resp.content)

                    if ext in (".xlsx", ".xls"):
                        # Reuse existing Excel loader
                        excel_engine_name = "openpyxl" if ext == ".xlsx" else "xlrd"
                        excel_file = pd.ExcelFile(temp_path, engine=excel_engine_name)
                        sheets_to_import = excel_file.sheet_names

                        if body.worksheet_name:
                            if body.worksheet_name in sheets_to_import:
                                sheets_to_import = [body.worksheet_name]
                            else:
                                raise HTTPException(
                                    status_code=400,
                                    detail=f"Worksheet '{body.worksheet_name}' not found. Available: {', '.join(sheets_to_import)}",
                                )

                        for sheet_name in sheets_to_import:
                            if len(sheets_to_import) > 1:
                                tname = f"{base_table}_{sanitize_table_name(sheet_name)}"
                            else:
                                tname = base_table

                            result = load_excel_sheet_to_postgres(
                                temp_path, sheet_name, tname,
                                has_header=True, engine=sql_engine,
                            )
                            if result["row_count"] == 0:
                                continue
                            total_rows += result["row_count"]
                            sheet_stats.append({
                                "sheet_name": sheet_name,
                                "table_name": result["table_name"],
                                "row_count": result["row_count"],
                                "column_count": result["column_count"],
                            })
                    else:
                        # Reuse existing CSV loader
                        result = load_csv_to_postgres(
                            temp_path, base_table,
                            delimiter=",", has_header=True, engine=sql_engine,
                        )
                        if result["row_count"] > 0:
                            total_rows = result["row_count"]
                            sheet_stats.append({
                                "sheet_name": file_name,
                                "table_name": result["table_name"],
                                "row_count": result["row_count"],
                                "column_count": result["column_count"],
                            })
                finally:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)

            else:
                # Native Google Sheet — use gspread
                if body.worksheet_name:
                    try:
                        worksheets = [spreadsheet.worksheet(body.worksheet_name)]
                    except Exception as e:
                        raise HTTPException(status_code=400, detail=f"Worksheet '{body.worksheet_name}' not found: {e}")
                else:
                    worksheets = spreadsheet.worksheets()

                if not worksheets:
                    raise HTTPException(status_code=400, detail="Spreadsheet has no worksheets")

                for ws in worksheets:
                    all_values = ws.get_all_values()
                    if not all_values or len(all_values) < 2:
                        continue

                    headers = all_values[0]
                    data_rows = all_values[1:]
                    df = pd.DataFrame(data_rows, columns=headers)

                    if len(worksheets) > 1:
                        table_name = f"{base_table}_{sanitize_table_name(ws.title)}"
                    else:
                        table_name = base_table

                    result = load_dataframe_to_postgres(
                        df, table_name, has_header=True, engine=sql_engine,
                    )

                    if result["row_count"] == 0:
                        continue

                    total_rows += result["row_count"]
                    sheet_stats.append({
                        "sheet_name": ws.title,
                        "table_name": result["table_name"],
                        "row_count": result["row_count"],
                        "column_count": result["column_count"],
                    })

            if not sheet_stats:
                raise HTTPException(status_code=400, detail="All worksheets are empty")

            # Create ONE connection pointing to the upload database (same as CSV)
            conn_id = str(uuid.uuid4())
            db_create_connection({
                "id": conn_id,
                "name": body.name,
                "db_type": "postgresql",
                "host": db_config["host"],
                "port": db_config["port"],
                "database": db_config["database"],
                "username": db_config["username"],
                "password": db_config["password"],
                "ssl_enabled": False,
            })
            new_conn = DatabaseConnectionConfig(
                id=conn_id,
                name=body.name,
                db_type=DatabaseType.POSTGRESQL,
                host=db_config["host"],
                port=db_config["port"],
                database=db_config["database"],
                username=db_config["username"],
                password=SecretStr(db_config["password"]),
                ssl_enabled=False,
                created_at=now,
                updated_at=now,
            )
            config.database_connections.append(new_conn)

            return JSONResponse(
                status_code=201,
                content={
                    "success": True,
                    "message": f"Google Sheet loaded into PostgreSQL: {len(sheet_stats)} sheet(s), {total_rows} total rows",
                    "connection_id": conn_id,
                    "name": body.name,
                    "db_type": "google_sheets",
                    "source_type": "google_sheets",
                    "host": db_config["host"],
                    "port": db_config["port"],
                    "database": db_config["database"],
                    "upload_tables": [s["table_name"] for s in sheet_stats],
                    "row_count": total_rows,
                    "sheet_stats": sheet_stats,
                },
            )

        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to load Google Sheet into database: {str(e)}")

    # ==========================================================================
    # Schema Sync
    # ==========================================================================

    @app.get("/api/v1/schema/sync", tags=["Schema"])
    async def sync_schema(
        connection_id: str,
        include_samples: bool = True,
        sample_limit: int = 10,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Sync schema from database connection.

        Returns schema with tables, columns, data types, and optional sample data.
        Returns database schema metadata with tables, columns, and optional sample data.
        """
        from sandbox.connectors.factory import get_connector
        from sandbox.core.config import DatabaseType, get_config

        config = get_config()

        # Find connection config
        conn_config = None
        for conn in config.database_connections:
            if conn.id == connection_id:
                conn_config = conn
                break

        if not conn_config:
            raise HTTPException(
                status_code=404,
                detail=f"Connection '{connection_id}' not found"
            )

        try:
            # Get connector
            connector = get_connector(conn_config.db_type, conn_config)

            # Get selected tables config (if any) — uploaded connections restrict
            # which tables from the shared schema belong to this connection
            selected_tables_config = conn_config.selected_tables or {}

            async with connector.get_connection() as conn:
                # Get all tables
                tables = await connector.get_tables(conn, schema=conn_config.schema_name)

                schema_data = {
                    "connection_id": connection_id,
                    "connection_name": conn_config.name,
                    "database": conn_config.database,
                    "db_type": conn_config.db_type.value,
                    "schema": conn_config.schema_name,
                    "tables": []
                }

                default_schema = "dbo" if conn_config.db_type == DatabaseType.MSSQL else "public"
                schema_prefix = conn_config.schema_name or default_schema

                # Batch-fetch all columns in one query if connector supports it
                all_columns_batch = None
                if hasattr(connector, 'get_all_columns'):
                    try:
                        all_columns_batch = await connector.get_all_columns(
                            conn, schema=conn_config.schema_name
                        )
                    except Exception as e:
                        logger.warning("batch_columns_failed", error=str(e))

                # Build tables list
                tables_to_process = []
                for table_name in tables:
                    if selected_tables_config:
                        full_name = f"{schema_prefix}.{table_name}"
                        table_selection = selected_tables_config.get(full_name)
                        if not table_selection or not table_selection.get("selected"):
                            continue
                        selected_columns = table_selection.get("columns", [])
                    else:
                        selected_columns = None
                    tables_to_process.append((table_name, selected_columns))

                for table_name, selected_columns in tables_to_process:
                    # Use batch result or fall back to per-table fetch
                    if all_columns_batch and table_name in all_columns_batch:
                        columns_info = all_columns_batch[table_name]
                    else:
                        columns_info = await connector.get_columns(
                            conn, table_name, schema=conn_config.schema_name
                        )

                    if selected_columns is not None and selected_columns:
                        columns_info = [
                            col for col in columns_info
                            if col.get("name") in selected_columns
                        ]

                    table_data = {
                        "name": table_name,
                        "columns": columns_info,
                        "sample_data": None
                    }

                    if include_samples:
                        try:
                            if selected_columns:
                                col_list = ", ".join(f'"{c}"' for c in selected_columns)
                            else:
                                col_list = "*"

                            is_mssql = conn_config.db_type == DatabaseType.MSSQL
                            top_clause = f"TOP {sample_limit} " if is_mssql else ""
                            limit_clause = "" if is_mssql else f" LIMIT {sample_limit}"

                            if conn_config.schema_name:
                                sample_query = f'SELECT {top_clause}{col_list} FROM "{conn_config.schema_name}"."{table_name}"{limit_clause}'
                            else:
                                sample_query = f'SELECT {top_clause}{col_list} FROM "{table_name}"{limit_clause}'

                            result = await connector.execute(conn, sample_query)

                            table_data["sample_data"] = {
                                "columns": result.columns,
                                "rows": [
                                    {col: _make_json_safe(val) for col, val in zip(result.columns, row)}
                                    for row in result.rows
                                ],
                                "total_rows": result.row_count
                            }
                        except Exception as e:
                            logger.warning("failed_to_get_samples", table=table_name, error=str(e))
                            table_data["sample_data"] = None

                    schema_data["tables"].append(table_data)

                return JSONResponse(content={
                    "status": "success",
                    "data": schema_data
                })

        except Exception as e:
            logger.error("schema_sync_error", connection_id=connection_id, error=str(e))
            raise HTTPException(status_code=500, detail=str(e))

    @app.get("/api/v1/schema/full-sync", tags=["Schema"])
    async def full_sync_schema(
        include_samples: bool = True,
        sample_limit: int = 10,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Bulk sync: returns all connections with schemas and sample data.

        Pulls all connection metadata in one call.
        No credentials are included in the response.
        """
        from sandbox.connectors.factory import get_connector
        from sandbox.core.config import DatabaseType, get_config

        config = get_config()

        async def _sync_table(connector, conn_config, table_name, selected_columns, include_samples, sample_limit):
            """Fetch columns and sample data for a single table."""
            async with connector.get_connection() as conn:
                columns_info = await connector.get_columns(
                    conn, table_name, schema=conn_config.schema_name
                )

                if selected_columns is not None and selected_columns:
                    columns_info = [
                        col for col in columns_info
                        if col.get("name") in selected_columns
                    ]

                table_data = {
                    "name": table_name,
                    "columns": columns_info,
                    "sample_data": None,
                }

                if include_samples:
                    try:
                        is_mssql = conn_config.db_type == DatabaseType.MSSQL

                        if is_mssql:
                            # T-SQL uses [schema].[table], not PostgreSQL-style double quotes.
                            def _mssql_esc(n: str) -> str:
                                return str(n).replace("]", "]]")

                            if selected_columns:
                                col_list = ", ".join(
                                    f"[{_mssql_esc(c)}]" for c in selected_columns
                                )
                            else:
                                col_list = "*"
                            top_clause = f"TOP {sample_limit} "
                            if conn_config.schema_name:
                                from_part = (
                                    f"[{_mssql_esc(conn_config.schema_name)}]"
                                    f".[{_mssql_esc(table_name)}]"
                                )
                            else:
                                from_part = f"[{_mssql_esc(table_name)}]"
                            sample_query = f"SELECT {top_clause}{col_list} FROM {from_part}"
                        else:
                            if selected_columns:
                                col_list = ", ".join(f'"{c}"' for c in selected_columns)
                            else:
                                col_list = "*"
                            top_clause = ""
                            limit_clause = f" LIMIT {sample_limit}"
                            if conn_config.schema_name:
                                sample_query = (
                                    f'SELECT {top_clause}{col_list} FROM '
                                    f'"{conn_config.schema_name}"."{table_name}"{limit_clause}'
                                )
                            else:
                                sample_query = (
                                    f'SELECT {top_clause}{col_list} FROM '
                                    f'"{table_name}"{limit_clause}'
                                )

                        result = await connector.execute(conn, sample_query)
                        table_data["sample_data"] = {
                            "columns": result.columns,
                            "rows": [
                                {col: _make_json_safe(val) for col, val in zip(result.columns, row)}
                                for row in result.rows
                            ],
                            "total_rows": result.row_count,
                        }
                    except Exception as e:
                        logger.warning(
                            "full_sync_sample_error",
                            connection=conn_config.id,
                            table=table_name,
                            error=str(e),
                        )

                return table_data

        async def _sync_connection(conn_config):
            """Sync schema for a single database connection."""
            is_upload = (
                conn_config.database.startswith("upload_")
                or conn_config.schema_name == "uploads"
            )
            display_db_type = "csv" if is_upload else conn_config.db_type.value
            connection_data = {
                "id": conn_config.id,
                "name": conn_config.name,
                "db_type": display_db_type,
                "host": conn_config.host,
                "port": conn_config.port,
                "database": conn_config.database,
                "schema": conn_config.schema_name,
                "is_default": getattr(conn_config, "is_default", False),
                "tables": [],
            }

            selected_tables_config = conn_config.selected_tables or {}

            try:
                connector = get_connector(conn_config.db_type, conn_config)

                async with connector.get_connection() as conn:
                    tables = await connector.get_tables(
                        conn, schema=conn_config.schema_name
                    )

                default_schema = "dbo" if conn_config.db_type == DatabaseType.MSSQL else "public"
                schema_prefix = conn_config.schema_name or default_schema

                # Build list of tables to sync with their column selections
                tables_to_sync = []
                for table_name in tables:
                    if selected_tables_config:
                        full_name = f"{schema_prefix}.{table_name}"
                        table_selection = selected_tables_config.get(full_name)
                        if not table_selection or not table_selection.get("selected"):
                            continue
                        selected_columns = table_selection.get("columns", [])
                    else:
                        selected_columns = None
                    tables_to_sync.append((table_name, selected_columns))

                # Try batch column fetch (1 query for ALL tables) if connector supports it
                all_columns_batch = None
                if tables_to_sync and hasattr(connector, 'get_all_columns'):
                    try:
                        async with connector.get_connection() as conn:
                            all_columns_batch = await connector.get_all_columns(
                                conn, schema=conn_config.schema_name
                            )
                        logger.info(
                            "full_sync_batch_columns",
                            connection=conn_config.id,
                            tables=len(all_columns_batch),
                        )
                    except Exception as e:
                        logger.warning(
                            "full_sync_batch_columns_failed",
                            connection=conn_config.id,
                            error=str(e),
                        )

                # Build table data
                if tables_to_sync:
                    if all_columns_batch:
                        # Fast path: use batch-fetched columns, only fetch sample data in parallel
                        async def _build_table_from_batch(tname, sel_cols):
                            columns_info = all_columns_batch.get(tname, [])
                            if sel_cols is not None and sel_cols:
                                columns_info = [c for c in columns_info if c.get("name") in sel_cols]

                            table_data = {"name": tname, "columns": columns_info, "sample_data": None}

                            if include_samples:
                                try:
                                    is_mssql = conn_config.db_type == DatabaseType.MSSQL

                                    def _mssql_esc2(n: str) -> str:
                                        return str(n).replace("]", "]]")

                                    if is_mssql:
                                        if sel_cols:
                                            col_list = ", ".join(
                                                f"[{_mssql_esc2(c)}]" for c in sel_cols
                                            )
                                        else:
                                            col_list = "*"
                                        top_clause = f"TOP {sample_limit} "
                                        if conn_config.schema_name:
                                            from_part = (
                                                f"[{_mssql_esc2(conn_config.schema_name)}]"
                                                f".[{_mssql_esc2(tname)}]"
                                            )
                                        else:
                                            from_part = f"[{_mssql_esc2(tname)}]"
                                        sample_query = (
                                            f"SELECT {top_clause}{col_list} FROM {from_part}"
                                        )
                                    else:
                                        if sel_cols:
                                            col_list = ", ".join(f'"{c}"' for c in sel_cols)
                                        else:
                                            col_list = "*"
                                        top_clause = ""
                                        limit_clause = f" LIMIT {sample_limit}"
                                        if conn_config.schema_name:
                                            sample_query = (
                                                f'SELECT {top_clause}{col_list} FROM '
                                                f'"{conn_config.schema_name}"."{tname}"{limit_clause}'
                                            )
                                        else:
                                            sample_query = (
                                                f'SELECT {top_clause}{col_list} FROM '
                                                f'"{tname}"{limit_clause}'
                                            )
                                    async with connector.get_connection() as sconn:
                                        result = await connector.execute(sconn, sample_query)
                                    table_data["sample_data"] = {
                                        "columns": result.columns,
                                        "rows": [
                                            {col: _make_json_safe(val) for col, val in zip(result.columns, row)}
                                            for row in result.rows
                                        ],
                                        "total_rows": result.row_count,
                                    }
                                except Exception as e:
                                    logger.warning("full_sync_sample_error", connection=conn_config.id, table=tname, error=str(e))

                            return table_data

                        table_results = await asyncio.gather(
                            *[_build_table_from_batch(tname, sel_cols) for tname, sel_cols in tables_to_sync],
                            return_exceptions=True,
                        )
                    else:
                        # Fallback: fetch columns per table in parallel
                        table_results = await asyncio.gather(
                            *[
                                _sync_table(connector, conn_config, tname, sel_cols, include_samples, sample_limit)
                                for tname, sel_cols in tables_to_sync
                            ],
                            return_exceptions=True,
                        )

                    for i, result in enumerate(table_results):
                        if isinstance(result, Exception):
                            logger.warning(
                                "full_sync_table_error",
                                connection=conn_config.id,
                                table=tables_to_sync[i][0],
                                error=str(result),
                            )
                        else:
                            connection_data["tables"].append(result)

            except Exception as e:
                logger.warning(
                    "full_sync_connection_error",
                    connection=conn_config.id,
                    error=str(e),
                )
                connection_data["error"] = str(e)

            return connection_data

        # Sync all connections in parallel
        results = await asyncio.gather(
            *[_sync_connection(cc) for cc in config.database_connections],
            return_exceptions=True,
        )
        synced_connections = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                cc = config.database_connections[i]
                logger.warning(
                    "full_sync_connection_error",
                    connection=cc.id,
                    error=str(result),
                )
                synced_connections.append({
                    "id": cc.id,
                    "name": cc.name,
                    "db_type": cc.db_type.value,
                    "tables": [],
                    "error": str(result),
                })
            else:
                synced_connections.append(result)

        return JSONResponse(content={
            "status": "success",
            "synced_at": datetime.now(timezone.utc).isoformat(),
            "connections": synced_connections,
        })

    @app.get("/api/v1/schema/table/{table_name}/samples", tags=["Schema"])
    async def get_table_samples(
        connection_id: str,
        table_name: str,
        limit: int = 10,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Get sample data from a specific table.
        """
        from sandbox.connectors.factory import get_connector
        from sandbox.core.config import get_config

        config = get_config()

        # Find connection config
        conn_config = None
        for conn in config.database_connections:
            if conn.id == connection_id:
                conn_config = conn
                break

        if not conn_config:
            raise HTTPException(
                status_code=404,
                detail=f"Connection '{connection_id}' not found"
            )

        try:
            connector = get_connector(conn_config.db_type, conn_config)

            async with connector.get_connection() as conn:
                # Build query
                if conn_config.schema_name:
                    query = f'SELECT * FROM "{conn_config.schema_name}"."{table_name}" LIMIT {limit}'
                else:
                    query = f'SELECT * FROM "{table_name}" LIMIT {limit}'

                result = await connector.execute(conn, query)

                return JSONResponse(content={
                    "columns": result.columns,
                    "rows": [
                        {col: _make_json_safe(val) for col, val in zip(result.columns, row)}
                        for row in result.rows
                    ],
                    "total_rows": result.row_count
                })

        except Exception as e:
            logger.error(
                "get_table_samples_error",
                connection_id=connection_id,
                table=table_name,
                error=str(e)
            )
            raise HTTPException(status_code=500, detail=str(e))

    # ==========================================================================
    # SQL Pad Integration
    # ==========================================================================

    @app.post("/api/v1/sqlpad/connection", tags=["SQL Pad"])
    async def create_sqlpad_connection(
        connection_id: str,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Create or update SQL Pad connection for a database.

        This syncs the sandbox database connection to SQL Pad so users can
        explore and query their data using SQL Pad's UI.
        """
        from sandbox.services.sqlpad_service import get_sqlpad_service
        from sandbox.core.config import get_config

        config = get_config()

        # Find connection in config
        conn_config = next(
            (c for c in config.database_connections if c.id == connection_id),
            None
        )

        if not conn_config:
            raise HTTPException(
                status_code=404,
                detail=f"Connection {connection_id} not found"
            )

        try:
            sqlpad = get_sqlpad_service()

            result = await sqlpad.create_or_update_connection(
                connection_id=conn_config.id,
                name=conn_config.name,
                db_type=conn_config.db_type.value,
                host=conn_config.host,
                port=conn_config.port,
                database=conn_config.database,
                username=conn_config.username,
                password=conn_config.password.get_secret_value(),
                schema=conn_config.schema_name,
            )

            return JSONResponse(content={
                "status": "success",
                "data": {
                    "connection_id": result.get("id"),
                    "name": result.get("name"),
                    "driver": result.get("driver"),
                }
            })

        except Exception as e:
            logger.error("sqlpad_connection_error", error=str(e))
            raise HTTPException(status_code=500, detail=str(e))

    @app.delete("/api/v1/sqlpad/connection/{connection_id}", tags=["SQL Pad"])
    async def delete_sqlpad_connection(
        connection_id: str,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """Delete a SQL Pad connection."""
        from sandbox.services.sqlpad_service import get_sqlpad_service

        try:
            sqlpad = get_sqlpad_service()
            await sqlpad.delete_connection(connection_id)

            return JSONResponse(content={
                "status": "success",
                "message": "Connection deleted"
            })

        except Exception as e:
            logger.error("sqlpad_delete_error", error=str(e))
            raise HTTPException(status_code=500, detail=str(e))

    @app.get("/api/v1/sqlpad/connections", tags=["SQL Pad"])
    async def list_sqlpad_connections(
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """List all SQL Pad connections."""
        from sandbox.services.sqlpad_service import get_sqlpad_service

        try:
            sqlpad = get_sqlpad_service()
            connections = await sqlpad.list_connections()

            return JSONResponse(content={
                "status": "success",
                "data": connections
            })

        except Exception as e:
            logger.error("sqlpad_list_error", error=str(e))
            raise HTTPException(status_code=500, detail=str(e))

    @app.get("/api/v1/sqlpad/embed-url", tags=["SQL Pad"])
    async def get_sqlpad_embed_url(
        connection_id: str | None = None,
        token_data: dict = Depends(verify_sandbox_token),
    ) -> JSONResponse:
        """
        Get SQL Pad embed URL with authentication token.

        Use this URL in an iframe to embed SQL Pad in your UI.
        """
        from sandbox.services.sqlpad_service import get_sqlpad_service

        try:
            sqlpad = get_sqlpad_service()
            embed_url = await sqlpad.get_embed_url(connection_id)

            return JSONResponse(content={
                "status": "success",
                "data": {
                    "embed_url": embed_url,
                    "connection_id": connection_id,
                }
            })

        except Exception as e:
            logger.error("sqlpad_embed_url_error", error=str(e))
            raise HTTPException(status_code=500, detail=str(e))

    # ==========================================================================
    # AI Query Generation (proxies to MVP backend)
    # ==========================================================================

    async def _extract_api_key(
        x_api_key: str | None = Header(None, alias="X-API-Key"),
        authorization: str | None = Header(None),
    ) -> str | None:
        """Extract the raw API key from request headers."""
        if x_api_key:
            return x_api_key
        if authorization and authorization.startswith("Bearer "):
            return authorization[7:]
        return None

    @app.post("/api/v1/ai/generate-query", tags=["AI"])
    async def ai_generate_query(
        request: AIGenerateQueryRequest,
        token_data: dict = Depends(verify_sandbox_token),
        api_key: str | None = Depends(_extract_api_key),
    ) -> JSONResponse:
        """
        Generate SQL query from natural language using AI.

        Proxies the request to the MVP backend which runs the LangGraph agent.
        The sandbox just forwards connection_id + user_query.
        """
        import httpx

        config = get_config()

        # Derive MVP base URL from remote auth URL
        # e.g. "http://host.docker.internal:18000/api/v1/sandbox/validate-key"
        # -> "http://host.docker.internal:18000"
        remote_url = getattr(config.authentication, "remote_url", "")
        if "/api/" in remote_url:
            mvp_base_url = remote_url.split("/api/")[0]
        else:
            mvp_base_url = remote_url.rstrip("/")

        if not mvp_base_url:
            raise HTTPException(
                status_code=503,
                detail="AI query generation not available (MVP URL not configured)",
            )

        headers = {"Content-Type": "application/json"}
        if api_key:
            headers["X-API-Key"] = api_key

        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                response = await client.post(
                    f"{mvp_base_url}/api/v1/generate-sql",
                    json={
                        "connection_id": request.connection_id,
                        "user_query": request.user_query,
                    },
                    headers=headers,
                )

            try:
                content = response.json()
            except Exception:
                content = {"success": False, "error": f"MVP returned HTTP {response.status_code}: {response.text[:200]}"}

            return JSONResponse(
                status_code=response.status_code if response.status_code < 500 else 502,
                content=content,
            )

        except httpx.TimeoutException:
            raise HTTPException(
                status_code=504,
                detail="AI query generation timed out",
            )
        except Exception as e:
            logger.error("ai_generate_query_error", error=str(e))
            raise HTTPException(status_code=502, detail=str(e))

    # ==========================================================================
    # Document Knowledge Base APIs
    # ==========================================================================
    # These endpoints store document files, KB vectors, and chunks in the
    # sandbox's own PostgreSQL (with pgvector). This keeps user data on their
    # infrastructure for security.

    def _get_doc_db_engine():
        """Reuse the sandbox's upload DB engine for document KB storage."""
        return _get_api_key_engine()

    def _ensure_doc_kb_tables():
        """Create document KB tables with pgvector in sandbox postgres. Idempotent."""
        from sqlalchemy import text as sql_text
        engine = _get_doc_db_engine()
        with engine.connect() as conn:
            try:
                conn.execute(sql_text("CREATE EXTENSION IF NOT EXISTS vector"))
                conn.commit()
            except Exception:
                conn.rollback()

            conn.execute(sql_text("""
                CREATE TABLE IF NOT EXISTS document_knowledge_bases (
                    id SERIAL PRIMARY KEY,
                    workspace_id TEXT,
                    connection_id TEXT,
                    name VARCHAR(200) NOT NULL,
                    description TEXT,
                    created_at TIMESTAMPTZ DEFAULT NOW(),
                    updated_at TIMESTAMPTZ DEFAULT NOW()
                )
            """))

            conn.execute(sql_text("""
                CREATE TABLE IF NOT EXISTS document_kb_documents (
                    id SERIAL PRIMARY KEY,
                    knowledge_base_id INTEGER NOT NULL REFERENCES document_knowledge_bases(id) ON DELETE CASCADE,
                    filename VARCHAR(500) NOT NULL,
                    file_type VARCHAR(20) NOT NULL,
                    file_size INTEGER,
                    storage_path TEXT,
                    source_path TEXT,
                    source_type VARCHAR(50),
                    status VARCHAR(20) DEFAULT 'pending',
                    error_message TEXT,
                    extracted_text TEXT,
                    summary TEXT,
                    chunk_count INTEGER DEFAULT 0,
                    progress INTEGER DEFAULT 0,
                    created_at TIMESTAMPTZ DEFAULT NOW(),
                    updated_at TIMESTAMPTZ DEFAULT NOW()
                )
            """))

            # Migration: replace file_data BYTEA with storage_path
            for col_sql in [
                "ALTER TABLE document_kb_documents ADD COLUMN IF NOT EXISTS storage_path TEXT",
                "ALTER TABLE document_kb_documents DROP COLUMN IF EXISTS file_data",
            ]:
                try:
                    conn.execute(sql_text(col_sql))
                except Exception:
                    pass

            # Migration: add path columns if table already exists
            for col_sql in [
                "ALTER TABLE document_kb_documents ADD COLUMN IF NOT EXISTS source_path TEXT",
                "ALTER TABLE document_kb_documents ADD COLUMN IF NOT EXISTS source_type VARCHAR(50)",
            ]:
                try:
                    conn.execute(sql_text(col_sql))
                except Exception:
                    pass

            conn.execute(sql_text("""
                CREATE TABLE IF NOT EXISTS document_kb_chunks (
                    id SERIAL PRIMARY KEY,
                    document_id INTEGER NOT NULL REFERENCES document_kb_documents(id) ON DELETE CASCADE,
                    knowledge_base_id INTEGER NOT NULL REFERENCES document_knowledge_bases(id) ON DELETE CASCADE,
                    chunk_index INTEGER NOT NULL,
                    content TEXT NOT NULL,
                    token_count INTEGER,
                    embedding vector(1536),
                    page_num INTEGER,
                    word_positions JSONB,
                    created_at TIMESTAMPTZ DEFAULT NOW()
                )
            """))

            # Indexes
            for idx_sql in [
                "CREATE INDEX IF NOT EXISTS idx_dkb_ws ON document_knowledge_bases(workspace_id)",
                "CREATE INDEX IF NOT EXISTS idx_dkbd_kb ON document_kb_documents(knowledge_base_id)",
                "CREATE INDEX IF NOT EXISTS idx_dkbd_status ON document_kb_documents(status)",
                "CREATE INDEX IF NOT EXISTS idx_dkbc_doc ON document_kb_chunks(document_id)",
                "CREATE INDEX IF NOT EXISTS idx_dkbc_kb ON document_kb_chunks(knowledge_base_id)",
            ]:
                conn.execute(sql_text(idx_sql))

            try:
                conn.execute(sql_text("""
                    CREATE INDEX IF NOT EXISTS idx_dkbc_embedding
                    ON document_kb_chunks USING ivfflat (embedding vector_cosine_ops) WITH (lists = 100)
                """))
            except Exception:
                pass

            conn.commit()
        logger.info("document_kb_tables_ensured")

    # Initialize tables on startup
    try:
        _ensure_doc_kb_tables()
    except Exception as e:
        logger.warning("document_kb_tables_init_failed", error=str(e))

    # --- Create Knowledge Base ---
    @app.post("/api/v1/documents/create-kb", tags=["Documents"])
    async def create_knowledge_base(request: Request):
        """Create a new knowledge base."""
        from sqlalchemy import text as sql_text
        body = await request.json()
        name = body.get("name", "Untitled KB")
        description = body.get("description", "")
        workspace_id = body.get("workspace_id", "")
        connection_id = body.get("connection_id", "")

        engine = _get_doc_db_engine()
        with engine.connect() as conn:
            # Check if already exists for this connection
            if connection_id:
                row = conn.execute(
                    sql_text("SELECT id FROM document_knowledge_bases WHERE connection_id = :cid"),
                    {"cid": connection_id},
                ).fetchone()
                if row:
                    return {"kb_id": str(row[0]), "name": name, "status": "exists"}

            conn.execute(
                sql_text("""
                    INSERT INTO document_knowledge_bases (workspace_id, connection_id, name, description)
                    VALUES (:ws, :cid, :name, :desc)
                """),
                {"ws": workspace_id, "cid": connection_id, "name": name, "desc": description},
            )
            conn.commit()
            row = conn.execute(
                sql_text("SELECT id FROM document_knowledge_bases WHERE connection_id = :cid ORDER BY id DESC LIMIT 1"),
                {"cid": connection_id or f"auto-{uuid.uuid4()}"},
            ).fetchone()
            kb_id = str(row[0]) if row else "0"

        return {"kb_id": kb_id, "name": name, "status": "created"}

    # --- Upload Document ---
    @app.post("/api/v1/documents/upload", tags=["Documents"])
    async def upload_document(
        file: UploadFile = File(...),
        workspace_id: str = Form(""),
        folder_path: str = Form(""),
    ):
        """
        Upload a document file to the sandbox filesystem.
        Preserves directory structure when folder_path is provided.

        Files are stored at: /app/data/documents/{workspace_id}/{folder_path}/{filename}
        No S3, no BYTEA — just local filesystem in the sandbox volume.
        """
        content = await file.read()
        file_id = str(uuid.uuid4())
        original_filename = file.filename or "unknown"
        ext = Path(original_filename).suffix

        # Build storage path preserving directory structure
        base_dir = Path("/app/data/documents")
        if workspace_id:
            base_dir = base_dir / workspace_id
        if folder_path:
            # Sanitize folder_path to prevent path traversal
            safe_folder = Path(folder_path.replace("..", "").strip("/"))
            base_dir = base_dir / safe_folder

        base_dir.mkdir(parents=True, exist_ok=True)

        # Use file_id + original extension for uniqueness, but keep original name accessible
        storage_filename = f"{file_id}{ext}"
        file_path = base_dir / storage_filename
        file_path.write_bytes(content)

        # Full path relative to /app/data/documents for retrieval
        relative_path = str(file_path.relative_to(Path("/app/data/documents")))

        return {
            "file_id": file_id,
            "filename": original_filename,
            "size": len(content),
            "path": str(file_path),
            "storage_path": relative_path,
            "folder_path": folder_path,
        }

    # --- Process Document ---
    @app.post("/api/v1/documents/process", tags=["Documents"])
    async def process_document(request: Request):
        """
        Process a document using Unstructured (parsing) + OpenAI (embeddings).
        Supports OCR strategies: "local" (Tesseract) or "google_vision".
        """
        from sqlalchemy import text as sql_text

        body = await request.json()
        file_id = body.get("file_id", "")
        filename = body.get("filename", "unknown")
        file_type = body.get("file_type", "txt")
        kb_id = int(body.get("kb_id", 0))
        workspace_id = body.get("workspace_id", "")
        ocr_strategy = body.get("ocr_strategy", "local")  # "local" or "google_vision"
        google_vision_credentials = body.get("google_vision_credentials")  # JSON string
        source_path = body.get("source_path", "")  # e.g. "Google Drive:/Reports/Q3/financial.pdf"
        source_type = body.get("source_type", "")  # e.g. "google_drive", "s3", "upload"

        # Find the uploaded file — search recursively because files are stored
        # at /app/data/documents/{workspace_id}/{folder_path}/{file_id}.ext
        data_dir = Path("/app/data/documents")
        file_path = None
        for p in data_dir.rglob(f"{file_id}*"):
            if p.is_file():
                file_path = p
                break

        if not file_path or not file_path.exists():
            raise HTTPException(status_code=404, detail=f"File {file_id} not found")

        file_size = file_path.stat().st_size
        # Compute storage_path relative to /app/data/documents
        try:
            storage_rel = str(file_path.relative_to(Path("/app/data/documents")))
        except ValueError:
            storage_rel = str(file_path)

        engine = _get_doc_db_engine()
        with engine.connect() as conn:
            conn.execute(
                sql_text("""
                    INSERT INTO document_kb_documents
                        (knowledge_base_id, filename, file_type, file_size, storage_path, source_path, source_type, status, progress)
                    VALUES (:kb_id, :fname, :ftype, :fsize, :spath_disk, :spath, :stype, 'processing', 0)
                """),
                {"kb_id": kb_id, "fname": filename, "ftype": file_type, "fsize": file_size,
                 "spath_disk": storage_rel, "spath": source_path or None, "stype": source_type or None},
            )
            conn.commit()
            row = conn.execute(
                sql_text("SELECT id FROM document_kb_documents WHERE knowledge_base_id = :kb_id AND filename = :fname ORDER BY id DESC LIMIT 1"),
                {"kb_id": kb_id, "fname": filename},
            ).fetchone()
            doc_id = row[0]

        # Run the pipeline synchronously in a thread pool so the HTTP request
        # waits for actual completion. This way the backend knows when indexing
        # is truly done and can mark the document as ready with accurate chunk counts.
        import asyncio
        try:
            await asyncio.to_thread(
                _run_unstructured_pipeline,
                engine, doc_id, str(file_path), file_type, kb_id, filename,
                ocr_strategy, google_vision_credentials, source_path,
            )
        except Exception as e:
            logger.error("document_processing_failed", doc_id=doc_id, error=str(e))
            return {"doc_id": str(doc_id), "status": "failed", "error": str(e), "file_id": file_id}

        # Check final status from sandbox DB
        from sqlalchemy import text as sql_text_sync
        with engine.connect() as c:
            status_row = c.execute(
                sql_text_sync("SELECT status, chunk_count, error_message FROM document_kb_documents WHERE id = :id"),
                {"id": doc_id},
            ).fetchone()

        if status_row:
            return {
                "doc_id": str(doc_id),
                "status": status_row[0],
                "chunk_count": status_row[1] or 0,
                "error": status_row[2],
                "file_id": file_id,
            }
        return {"doc_id": str(doc_id), "status": "ready", "file_id": file_id}

    def _run_unstructured_pipeline(
        engine, doc_id: int, file_path: str, file_type: str,
        kb_id: int, filename: str,
        ocr_strategy: str = "local",
        google_vision_credentials: str | None = None,
        source_path: str = "",
    ):
        """
        Document processing pipeline using Unstructured (parsing) + LlamaIndex (indexing).

        - Unstructured: extracts structured elements (text, tables, headers, images/OCR)
        - LlamaIndex: chunks with metadata, embeds with OpenAI, stores in PGVectorStore

        OCR strategies for scanned/non-digital documents:
          - "local": Tesseract OCR (free, on-premise)
          - "google_vision": Google Cloud Vision OCR (best accuracy)

        LlamaIndex automatically tracks file paths, page numbers, and element types
        as metadata on each node, enabling path-based filtering during retrieval.
        """
        from sqlalchemy import text as sql_text
        import json as _json

        def _update_prog(prog):
            with engine.connect() as c:
                c.execute(sql_text("UPDATE document_kb_documents SET progress=:p, updated_at=NOW() WHERE id=:id"), {"p": prog, "id": doc_id})
                c.commit()

        AUDIO_TYPES = {"mp3", "wav", "m4a", "ogg", "flac", "wma"}
        VIDEO_TYPES = {"mp4", "mkv", "avi", "mov", "webm"}

        try:
            _update_prog(5)

            from llama_index.core.schema import Document as LIDocument
            li_documents = []
            extracted_text_parts = []
            is_media = file_type.lower() in (AUDIO_TYPES | VIDEO_TYPES)

            if is_media:
                # ── Audio/Video: Transcribe with faster-whisper ──
                logger.info("whisper_transcription_start", doc_id=doc_id, filename=filename, file_type=file_type)

                audio_path = file_path

                # If video, extract audio track to /tmp (not next to source — that breaks file lookups)
                if file_type.lower() in VIDEO_TYPES:
                    import subprocess
                    import tempfile
                    audio_fd, audio_path = tempfile.mkstemp(suffix=".wav", prefix=f"whisper_{doc_id}_")
                    os.close(audio_fd)  # We only need the path
                    subprocess.run(
                        ["ffmpeg", "-i", file_path, "-vn", "-acodec", "pcm_s16le",
                         "-ar", "16000", "-ac", "1", audio_path, "-y"],
                        capture_output=True, check=True, timeout=600,
                    )
                    logger.info("video_audio_extracted", doc_id=doc_id, audio_path=audio_path)

                _update_prog(10)

                # Transcribe with faster-whisper (small model, CPU)
                from faster_whisper import WhisperModel
                # Use pre-downloaded model to avoid runtime download + lock contention
                whisper_path = os.environ.get("WHISPER_MODEL_PATH", "")
                model = WhisperModel(
                    "small", device="cpu", compute_type="int8",
                    download_root=whisper_path or None,
                )
                segments, info = model.transcribe(audio_path, beam_size=5)

                logger.info("whisper_transcribing", doc_id=doc_id,
                            language=info.language, duration=f"{info.duration:.1f}s")

                _update_prog(15)

                # Build timestamped transcript as LlamaIndex documents
                segment_list = list(segments)
                for i, seg in enumerate(segment_list):
                    text = seg.text.strip()
                    if not text:
                        continue

                    # Format timestamp
                    start_m, start_s = divmod(int(seg.start), 60)
                    start_h, start_m = divmod(start_m, 60)
                    end_m, end_s = divmod(int(seg.end), 60)
                    end_h, end_m = divmod(end_m, 60)
                    timestamp = f"{start_h:02d}:{start_m:02d}:{start_s:02d} → {end_h:02d}:{end_m:02d}:{end_s:02d}"

                    extracted_text_parts.append(f"[{timestamp}] {text}")

                    metadata = {
                        "filename": filename,
                        "file_type": file_type,
                        "file_path": source_path or filename,
                        "source_path": source_path,
                        "element_type": "Transcript",
                        "doc_id": doc_id,
                        "kb_id": kb_id,
                        "timestamp_start": round(seg.start, 2),
                        "timestamp_end": round(seg.end, 2),
                        "timestamp": timestamp,
                        "language": info.language,
                    }
                    if source_path:
                        parts = source_path.rsplit("/", 1)
                        metadata["directory"] = parts[0] if len(parts) > 1 else ""

                    li_documents.append(LIDocument(text=f"[{timestamp}] {text}", metadata=metadata))

                    # Update progress proportionally
                    if i % 20 == 0:
                        pct = 15 + int((i / max(len(segment_list), 1)) * 5)
                        _update_prog(min(pct, 20))

                # Clean up the tmp audio file (extracted from video)
                # audio_path is in /tmp/, file_path is the original in /app/data/documents/
                if audio_path != file_path and os.path.exists(audio_path):
                    try:
                        os.remove(audio_path)
                        logger.info("temp_audio_cleaned", doc_id=doc_id)
                    except OSError:
                        pass

                logger.info("whisper_transcription_done", doc_id=doc_id,
                            segments=len(li_documents), language=info.language)

            else:
                # ── Documents: Parse with Unstructured ──
                logger.info("unstructured_parsing_start", doc_id=doc_id, filename=filename, ocr=ocr_strategy)

                from unstructured.partition.auto import partition

                partition_kwargs = {
                    "filename": file_path,
                    "strategy": "hi_res",
                }

                if ocr_strategy == "google_vision" and google_vision_credentials:
                    import tempfile
                    creds_file = tempfile.NamedTemporaryFile(suffix=".json", delete=False, mode="w")
                    creds_file.write(google_vision_credentials)
                    creds_file.close()
                    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = creds_file.name
                    os.environ["OCR_AGENT"] = "unstructured.partition.utils.ocr_models.google_vision_ocr.OCRAgentGoogleVision"
                else:
                    os.environ["OCR_AGENT"] = "unstructured.partition.utils.ocr_models.tesseract_ocr.OCRAgentTesseract"

                elements = partition(**partition_kwargs)

                if not elements:
                    raise ValueError("Unstructured could not extract any content from the document")

                # For PDFs: extract word-level bounding boxes using PyMuPDF
                # Unstructured handles parsing/OCR, PyMuPDF provides precise word positions
                pdf_page_words = {}  # page_num → list of {"text", "bbox", "page_width", "page_height"}
                if file_type.lower() == "pdf":
                    try:
                        import fitz
                        doc = fitz.open(file_path)
                        for page_idx in range(len(doc)):
                            page = doc[page_idx]
                            page_num_1 = page_idx + 1
                            words = []
                            for w in page.get_text("words"):
                                if w[4].strip():
                                    words.append({
                                        "text": w[4],
                                        "bbox": [round(w[0], 1), round(w[1], 1), round(w[2], 1), round(w[3], 1)],
                                    })
                            pdf_page_words[page_num_1] = {
                                "words": words,
                                "page_width": round(page.rect.width, 1),
                                "page_height": round(page.rect.height, 1),
                            }
                        doc.close()
                        logger.info("pymupdf_words_extracted", doc_id=doc_id, pages=len(pdf_page_words))
                    except Exception as e:
                        logger.warning("pymupdf_extraction_failed", doc_id=doc_id, error=str(e))

                # Track running character offset for highlight positioning
                char_offset = 0

                for el in elements:
                    text = str(el).strip()
                    if not text:
                        continue

                    el_type = type(el).__name__
                    page_num = el.metadata.page_number if hasattr(el.metadata, 'page_number') else None

                    # PPTX: page_number = slide number
                    # XLSX: page_name = sheet name
                    page_name = getattr(el.metadata, 'page_name', None)  # Sheet name for XLSX

                    # Extract bounding box coordinates from Unstructured (hi_res strategy)
                    bbox = None
                    page_width = None
                    page_height = None
                    coords = getattr(el.metadata, 'coordinates', None)
                    if coords and coords.points and coords.system:
                        # points is list of (x, y) tuples defining the bounding polygon
                        pts = coords.points
                        x_vals = [p[0] for p in pts]
                        y_vals = [p[1] for p in pts]
                        bbox = [round(min(x_vals), 1), round(min(y_vals), 1),
                                round(max(x_vals), 1), round(max(y_vals), 1)]
                        page_width = round(coords.system.width, 1) if hasattr(coords.system, 'width') else None
                        page_height = round(coords.system.height, 1) if hasattr(coords.system, 'height') else None

                    # Format special elements for better context
                    if el_type == "Table":
                        text = f"[TABLE]\n{text}\n[/TABLE]"
                    elif el_type == "Title":
                        text = f"## {text}"
                    elif el_type == "Image":
                        text = f"[IMAGE TEXT] {text} [/IMAGE TEXT]"

                    # Track start/end char positions in the full extracted_text
                    start_char = char_offset
                    end_char = char_offset + len(text)
                    char_offset = end_char + 2  # +2 for the "\n\n" separator

                    extracted_text_parts.append(text)

                    # Create LlamaIndex Document with rich metadata
                    metadata = {
                        "filename": filename,
                        "file_type": file_type,
                        "file_path": source_path or filename,
                        "source_path": source_path,
                        "source_type": "upload" if not source_path else source_path.split(":")[0] if ":" in source_path else "upload",
                        "element_type": el_type,
                        "doc_id": doc_id,
                        "kb_id": kb_id,
                        "start_char": start_char,
                        "end_char": end_char,
                    }
                    if page_num is not None:
                        metadata["page_num"] = page_num
                    if page_name:
                        metadata["page_name"] = page_name  # Sheet name (XLSX) or slide label (PPTX)

                    # For PDFs: map element text to word-level bounding boxes
                    if page_num and page_num in pdf_page_words:
                        page_data = pdf_page_words[page_num]
                        page_words_list = page_data["words"]
                        pw = page_data["page_width"]
                        ph = page_data["page_height"]
                        metadata["page_width"] = pw
                        metadata["page_height"] = ph

                        # Find matching words on this page for this element's text
                        el_text_clean = str(el).replace("\n", " ").strip()
                        el_first_words = el_text_clean.split()[:3]
                        page_word_texts = [w["text"] for w in page_words_list]

                        matched_bboxes = []
                        for si in range(len(page_word_texts)):
                            if page_word_texts[si:si + len(el_first_words)] == el_first_words:
                                # Found start — collect word bboxes for the element length
                                end_idx = min(si + len(el_text_clean.split()) + 5, len(page_words_list))
                                matched_bboxes = [{"bbox": w["bbox"]} for w in page_words_list[si:end_idx]]
                                break

                        if matched_bboxes:
                            import json as _json2
                            metadata["word_positions_json"] = _json2.dumps({
                                "page_num": page_num,
                                "page_width": pw,
                                "page_height": ph,
                                "word_count": len(matched_bboxes),
                                "words": matched_bboxes[:200],  # Limit to 200 words max
                            })

                    # Extract directory path from source_path for folder-level filtering
                    if source_path:
                        parts = source_path.rsplit("/", 1)
                        metadata["directory"] = parts[0] if len(parts) > 1 else ""

                    li_documents.append(LIDocument(text=text, metadata=metadata))

            extracted_text = "\n\n".join(extracted_text_parts)

            if not extracted_text.strip():
                raise ValueError("No text content extracted from document")

            # Save extracted text to DB
            with engine.connect() as c:
                c.execute(sql_text("UPDATE document_kb_documents SET extracted_text=:t, updated_at=NOW() WHERE id=:id"),
                          {"t": extracted_text, "id": doc_id})
                c.commit()

            logger.info("parsing_done", doc_id=doc_id,
                        li_docs=len(li_documents), chars=len(extracted_text),
                        media=is_media)
            _update_prog(30)

            # ── Step 3: Build LlamaIndex VectorStoreIndex with PGVectorStore ──
            from llama_index.core import VectorStoreIndex, Settings, StorageContext
            from llama_index.core.node_parser import SentenceSplitter
            from llama_index.embeddings.openai import OpenAIEmbedding
            from llama_index.vector_stores.postgres import PGVectorStore

            openai_key = os.environ.get("OPENAI_API_KEY") or os.environ.get("OPENAI_KEY", "")

            # Configure LlamaIndex settings
            Settings.embed_model = OpenAIEmbedding(
                model="text-embedding-3-small",
                dimensions=1536,
                api_key=openai_key,
            )
            Settings.chunk_size = 512
            Settings.chunk_overlap = 50

            # Connect LlamaIndex to sandbox postgres via PGVectorStore
            db_host = os.environ.get("SANDBOX_UPLOAD_DB_HOST", "sandbox-postgres")
            db_port = os.environ.get("SANDBOX_UPLOAD_DB_PORT", "5432")
            db_name = os.environ.get("SANDBOX_UPLOAD_DB_NAME", "sandbox_uploads")
            db_user = os.environ.get("SANDBOX_UPLOAD_DB_USER", "sandbox")
            db_pass = os.environ.get("SANDBOX_UPLOAD_DB_PASSWORD", "sandbox_password")

            # Each KB gets its own table for isolation
            table_name = f"llamaindex_kb_{kb_id}"

            vector_store = PGVectorStore.from_params(
                host=db_host,
                port=db_port,
                database=db_name,
                user=db_user,
                password=db_pass,
                table_name=table_name,
                embed_dim=1536,
            )

            # Wrap in StorageContext — this is what actually makes VectorStoreIndex
            # persist nodes to the vector store (without it, the index is in-memory only)
            storage_context = StorageContext.from_defaults(vector_store=vector_store)

            _update_prog(40)

            # ── Step 4: Index documents — LlamaIndex handles chunking + embedding + storage ──
            # SentenceSplitter respects element boundaries better than character splitting
            node_parser = SentenceSplitter(chunk_size=512, chunk_overlap=50)

            logger.info("llamaindex_indexing_start", doc_id=doc_id, documents=len(li_documents))

            index = VectorStoreIndex.from_documents(
                li_documents,
                storage_context=storage_context,
                transformations=[node_parser],
                show_progress=False,
            )

            # Count how many nodes (chunks) were actually created and persisted
            # Query the vector store directly to confirm persistence
            chunk_count = len(li_documents)  # Fallback estimate
            try:
                all_nodes = index.docstore.docs
                if all_nodes:
                    chunk_count = len(all_nodes)
            except Exception:
                pass

            logger.info("llamaindex_indexing_done", doc_id=doc_id, chunks=chunk_count, table=table_name)
            _update_prog(85)

            # ── Step 5: Generate summary ──
            import openai as _openai
            client = _openai.OpenAI(api_key=openai_key)
            summary = ""
            try:
                summary_resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Summarize this document in 2-3 sentences."},
                        {"role": "user", "content": extracted_text[:8000]},
                    ],
                    max_tokens=200,
                )
                summary = summary_resp.choices[0].message.content or ""
            except Exception:
                pass

            # ── Finalize ──
            with engine.connect() as c:
                c.execute(
                    sql_text("UPDATE document_kb_documents SET status='ready', chunk_count=:cc, summary=:s, progress=100, updated_at=NOW() WHERE id=:id"),
                    {"cc": chunk_count, "s": summary, "id": doc_id},
                )
                c.commit()

            logger.info("document_processed", doc_id=doc_id, filename=filename,
                        chunks=chunk_count, ocr=ocr_strategy, source_path=source_path)

        except Exception as e:
            import traceback
            tb = traceback.format_exc()
            logger.error("document_processing_error", doc_id=doc_id, error=str(e), traceback=tb)
            with engine.connect() as c:
                c.execute(
                    sql_text("UPDATE document_kb_documents SET status='failed', error_message=:err, updated_at=NOW() WHERE id=:id"),
                    {"err": f"{str(e)}\n{tb[-500:]}", "id": doc_id},
                )
                c.commit()

    def _json_parse_safe(val):
        """Parse a JSON string, return None if invalid."""
        if not val:
            return None
        if isinstance(val, dict):
            return val
        try:
            import json
            return json.loads(val)
        except Exception:
            return None

    # --- Query KB Vectors ---
    @app.post("/api/v1/documents/query", tags=["Documents"])
    async def query_kb_vectors(request: Request):
        """
        Search for similar document chunks using LlamaIndex's VectorStoreIndex.
        Supports path-based filtering (search within specific directories).
        """
        body = await request.json()
        query_embedding = body.get("query_embedding", [])
        kb_ids = body.get("kb_ids", [])
        top_k = body.get("top_k", 10)
        threshold = body.get("similarity_threshold", 0.3)
        # Optional path filter — search only within a specific directory
        path_filter = body.get("path_filter", "")  # e.g. "Google Drive:/Reports"

        if not query_embedding or not kb_ids:
            return {"chunks": []}

        from llama_index.core import VectorStoreIndex, Settings
        from llama_index.core.vector_stores.types import (
            VectorStoreQuery,
            VectorStoreQueryMode,
            MetadataFilters,
            MetadataFilter,
            FilterOperator,
        )
        from llama_index.vector_stores.postgres import PGVectorStore
        from llama_index.embeddings.openai import OpenAIEmbedding

        openai_key = os.environ.get("OPENAI_API_KEY") or os.environ.get("OPENAI_KEY", "")
        Settings.embed_model = OpenAIEmbedding(
            model="text-embedding-3-small", dimensions=1536, api_key=openai_key,
        )

        db_host = os.environ.get("SANDBOX_UPLOAD_DB_HOST", "sandbox-postgres")
        db_port = os.environ.get("SANDBOX_UPLOAD_DB_PORT", "5432")
        db_name = os.environ.get("SANDBOX_UPLOAD_DB_NAME", "sandbox_uploads")
        db_user = os.environ.get("SANDBOX_UPLOAD_DB_USER", "sandbox")
        db_pass = os.environ.get("SANDBOX_UPLOAD_DB_PASSWORD", "sandbox_password")

        all_chunks = []

        for kb_id in kb_ids:
            table_name = f"llamaindex_kb_{kb_id}"

            try:
                vector_store = PGVectorStore.from_params(
                    host=db_host, port=db_port, database=db_name,
                    user=db_user, password=db_pass,
                    table_name=table_name, embed_dim=1536,
                )

                # Build metadata filters
                filters = []
                if path_filter:
                    filters.append(MetadataFilter(
                        key="source_path", value=path_filter, operator=FilterOperator.CONTAINS,
                    ))

                # Query using LlamaIndex's vector store directly
                query = VectorStoreQuery(
                    query_embedding=query_embedding,
                    similarity_top_k=top_k,
                    mode=VectorStoreQueryMode.DEFAULT,
                    filters=MetadataFilters(filters=filters) if filters else None,
                )

                result = vector_store.query(query)

                for node, similarity in zip(result.nodes or [], result.similarities or []):
                    if similarity < threshold:
                        continue
                    meta = node.metadata or {}
                    all_chunks.append({
                        "chunk_id": node.node_id,
                        "document_id": str(meta.get("doc_id", "")),
                        "kb_id": str(kb_id),
                        "content": node.get_content(),
                        "similarity": float(similarity),
                        "page_num": meta.get("page_num"),
                        "word_positions": _json_parse_safe(meta.get("word_positions_json")),
                        "filename": meta.get("filename", ""),
                        "source_path": meta.get("source_path", ""),
                        "source_type": meta.get("source_type", ""),
                        "directory": meta.get("directory", ""),
                        "element_type": meta.get("element_type", ""),
                        # Text highlight offsets — from LlamaIndex node or element metadata
                        "start_char": node.start_char_idx if node.start_char_idx is not None else meta.get("start_char"),
                        "end_char": node.end_char_idx if node.end_char_idx is not None else meta.get("end_char"),
                        # PDF bounding box [x0, y0, x1, y1] in page coordinates
                        "bbox": meta.get("bbox"),
                        "page_width": meta.get("page_width"),
                        "page_height": meta.get("page_height"),
                        # Sheet name (XLSX) or slide label (PPTX)
                        "page_name": meta.get("page_name"),
                        # Audio/video timestamp offsets
                        "timestamp_start": meta.get("timestamp_start"),
                        "timestamp_end": meta.get("timestamp_end"),
                        "timestamp": meta.get("timestamp"),
                        "file_type": meta.get("file_type", ""),
                    })

            except Exception as e:
                logger.warning("kb_query_failed", kb_id=kb_id, error=str(e))

        # Sort by similarity descending, take top_k
        all_chunks.sort(key=lambda c: c["similarity"], reverse=True)
        all_chunks = all_chunks[:top_k]

        return {"chunks": all_chunks}

    # --- KB Status ---
    @app.get("/api/v1/documents/kb-status", tags=["Documents"])
    async def get_kb_status(kb_id: str, workspace_id: str = ""):
        """Get knowledge base processing status."""
        from sqlalchemy import text as sql_text

        engine = _get_doc_db_engine()
        with engine.connect() as conn:
            result = conn.execute(
                sql_text("""
                    SELECT
                        COUNT(*) as total,
                        COUNT(*) FILTER (WHERE status = 'ready') as ready,
                        COUNT(*) FILTER (WHERE status = 'processing') as processing,
                        COUNT(*) FILTER (WHERE status = 'failed') as failed,
                        COUNT(*) FILTER (WHERE status = 'pending') as pending,
                        COALESCE(SUM(chunk_count), 0) as total_chunks
                    FROM document_kb_documents
                    WHERE knowledge_base_id = :kb_id
                """),
                {"kb_id": int(kb_id)},
            )
            row = result.fetchone()
            if not row:
                return {"total": 0, "ready": 0, "processing": 0, "failed": 0, "pending": 0, "total_chunks": 0}

            return {
                "total": row[0], "ready": row[1], "processing": row[2],
                "failed": row[3], "pending": row[4], "total_chunks": row[5],
            }

    # --- List KB Documents ---
    @app.get("/api/v1/documents/list", tags=["Documents"])
    async def list_kb_documents(kb_id: str, workspace_id: str = ""):
        """List all documents in a knowledge base."""
        from sqlalchemy import text as sql_text

        engine = _get_doc_db_engine()
        with engine.connect() as conn:
            result = conn.execute(
                sql_text("""
                    SELECT id, filename, file_type, file_size, status, error_message,
                           chunk_count, progress, created_at, updated_at,
                           source_path, source_type, storage_path
                    FROM document_kb_documents
                    WHERE knowledge_base_id = :kb_id
                    ORDER BY created_at DESC
                """),
                {"kb_id": int(kb_id)},
            )
            docs = []
            for row in result.fetchall():
                docs.append({
                    "id": row[0], "filename": row[1], "file_type": row[2],
                    "file_size": row[3], "status": row[4], "error_message": row[5],
                    "chunk_count": row[6], "progress": row[7] or 0,
                    "created_at": row[8].isoformat() if row[8] else None,
                    "updated_at": row[9].isoformat() if row[9] else None,
                    "source_path": row[10], "source_type": row[11],
                    "storage_path": row[12],
                })

        return {"documents": docs}

    # Cache of transcoded videos so repeat views are instant
    _transcoded_video_cache: dict = {}

    def _get_video_codec(file_path: str) -> tuple:
        """Return (video_codec, audio_codec) using ffprobe. Empty strings on failure."""
        import subprocess
        try:
            result = subprocess.run(
                ["ffprobe", "-v", "error", "-select_streams", "v:0",
                 "-show_entries", "stream=codec_name",
                 "-of", "default=nokey=1:noprint_wrappers=1", str(file_path)],
                capture_output=True, text=True, timeout=10,
            )
            vcodec = result.stdout.strip()
            result = subprocess.run(
                ["ffprobe", "-v", "error", "-select_streams", "a:0",
                 "-show_entries", "stream=codec_name",
                 "-of", "default=nokey=1:noprint_wrappers=1", str(file_path)],
                capture_output=True, text=True, timeout=10,
            )
            acodec = result.stdout.strip()
            return vcodec, acodec
        except Exception:
            return "", ""

    def _transcode_video_to_h264(source_path: str) -> Path:
        """
        Transcode a video to H.264 + AAC for browser compatibility.
        Results are cached so repeat requests don't re-transcode.
        """
        import hashlib
        import subprocess
        src = Path(source_path)
        # Cache key based on source file path + mtime
        mtime = src.stat().st_mtime
        cache_key = hashlib.md5(f"{src}:{mtime}".encode()).hexdigest()
        if cache_key in _transcoded_video_cache:
            cached = _transcoded_video_cache[cache_key]
            if Path(cached).exists():
                return Path(cached)

        cache_dir = Path("/tmp/sandbox_video_cache")
        cache_dir.mkdir(parents=True, exist_ok=True)
        out_path = cache_dir / f"{cache_key}.mp4"

        logger.info("transcoding_video", source=str(src), target=str(out_path))
        try:
            # -c:v libx264: H.264 video (universal browser support)
            # -c:a aac: AAC audio (universal)
            # -preset ultrafast: fastest encode, slightly larger file
            # -movflags +faststart: put metadata at start for streaming
            # -y: overwrite output
            subprocess.run(
                ["ffmpeg", "-i", str(src),
                 "-c:v", "libx264", "-preset", "ultrafast", "-crf", "23",
                 "-c:a", "aac", "-b:a", "128k",
                 "-movflags", "+faststart",
                 "-y", str(out_path)],
                capture_output=True, check=True, timeout=600,
            )
            _transcoded_video_cache[cache_key] = str(out_path)
            return out_path
        except subprocess.TimeoutExpired:
            logger.error("video_transcode_timeout", source=str(src))
            return src
        except subprocess.CalledProcessError as e:
            logger.error("video_transcode_failed", error=str(e.stderr[:500]) if e.stderr else "unknown")
            return src

    def _serve_file_with_viewer_conversion(file_path, original_filename: str = ""):
        """
        Serve a file, converting DOCX/XLSX to browser-viewable HTML.
        For videos, transcode to H.264 if the codec isn't browser-compatible.
        """
        ext = file_path.suffix.lower() if hasattr(file_path, 'suffix') else Path(file_path).suffix.lower()

        # Video: transcode to H.264 if needed
        if ext in (".mp4", ".mkv", ".avi", ".mov", ".webm"):
            vcodec, acodec = _get_video_codec(str(file_path))
            # H.264 and VP8/VP9 are browser-safe; others need transcoding
            browser_safe_video = vcodec in ("h264", "vp8", "vp9", "av1")
            browser_safe_audio = acodec in ("aac", "mp3", "opus", "vorbis", "")
            if not (browser_safe_video and browser_safe_audio):
                logger.info("video_needs_transcoding", vcodec=vcodec, acodec=acodec)
                file_path = _transcode_video_to_h264(str(file_path))
                ext = ".mp4"

        # DOCX → HTML conversion (browsers can't render .docx inline)
        if ext == ".docx":
            try:
                import docx as _docx
                doc = _docx.Document(str(file_path))
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        paragraphs.append("<br>")
                        continue
                    # Map Word heading styles to HTML
                    style = (para.style.name or "").lower() if para.style else ""
                    if "heading 1" in style:
                        paragraphs.append(f"<h1>{text}</h1>")
                    elif "heading 2" in style:
                        paragraphs.append(f"<h2>{text}</h2>")
                    elif "heading 3" in style:
                        paragraphs.append(f"<h3>{text}</h3>")
                    elif "title" in style:
                        paragraphs.append(f"<h1 class='title'>{text}</h1>")
                    else:
                        paragraphs.append(f"<p>{text}</p>")

                # Extract tables too
                for table in doc.tables:
                    rows_html = []
                    for row in table.rows:
                        cells = "".join(f"<td>{cell.text.strip()}</td>" for cell in row.cells)
                        rows_html.append(f"<tr>{cells}</tr>")
                    if rows_html:
                        paragraphs.append(f"<table>{''.join(rows_html)}</table>")

                html = (
                    "<!DOCTYPE html><html><head>"
                    '<meta charset="utf-8">'
                    f'<title>{original_filename or file_path.name}</title>'
                    "<style>"
                    "body{font-family:system-ui,-apple-system,sans-serif;max-width:900px;margin:40px auto;padding:0 24px;line-height:1.7;color:#1e293b;background:#fff}"
                    "@media(prefers-color-scheme:dark){body{color:#e2e8f0;background:#1a1a2e}}"
                    "h1,h2,h3{color:inherit;margin-top:1.5em;margin-bottom:0.5em}"
                    "h1.title{font-size:2em;text-align:center;margin-top:0.5em}"
                    "p{margin:0.8em 0;text-align:justify}"
                    "table{border-collapse:collapse;margin:1em 0;width:100%}"
                    "td,th{border:1px solid #cbd5e1;padding:8px 12px;text-align:left}"
                    "@media(prefers-color-scheme:dark){td,th{border-color:#475569}}"
                    "</style></head><body>"
                    + "\n".join(paragraphs) + "</body></html>"
                )
                return Response(
                    content=html,
                    media_type="text/html; charset=utf-8",
                    headers={"Content-Disposition": "inline"},
                )
            except Exception as e:
                logger.warning("docx_conversion_failed", error=str(e))
                # Fall through to raw download

        # Default: serve with correct MIME
        ct_map = {
            ".pdf": "application/pdf",
            ".txt": "text/plain; charset=utf-8",
            ".md": "text/plain; charset=utf-8",
            ".csv": "text/csv; charset=utf-8",
            ".html": "text/html; charset=utf-8",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".webp": "image/webp",
            ".svg": "image/svg+xml",
            ".mp3": "audio/mpeg",
            ".wav": "audio/wav",
            ".m4a": "audio/mp4",
            ".ogg": "audio/ogg",
            ".flac": "audio/flac",
            ".mp4": "video/mp4",
            ".webm": "video/webm",
            ".mkv": "video/x-matroska",
            ".mov": "video/quicktime",
        }
        return Response(
            content=Path(file_path).read_bytes(),
            media_type=ct_map.get(ext, "application/octet-stream"),
            headers={"Content-Disposition": "inline"},
        )

    # --- Serve Document File ---
    @app.get("/api/v1/documents/file/{file_id}", tags=["Documents"])
    async def get_document_file(file_id: str, workspace_id: str = ""):
        """Serve a document file from sandbox storage."""
        # Find the file by matching the stem (filename without extensions)
        # Use strict matching: filename must START with file_id and have ONE extension after
        # This prevents matching side-files like "{uuid}.mp4.wav" when looking for "{uuid}.mp4"
        data_dir = Path("/app/data/documents")
        best_match = None
        for p in data_dir.rglob(f"{file_id}*"):
            if not p.is_file():
                continue
            # Strict: filename is exactly "{file_id}.{ext}" (no extra extensions like .mp4.wav)
            name_after_id = p.name[len(file_id):]
            if name_after_id.count(".") == 1:  # exactly one dot → single extension
                best_match = p
                break
            # Keep as fallback if nothing better found
            if best_match is None:
                best_match = p

        if best_match:
            return _serve_file_with_viewer_conversion(best_match)

        # Fall back to DB storage_path lookup (for files stored via storage_path column)
        from sqlalchemy import text as sql_text
        engine = _get_doc_db_engine()
        with engine.connect() as conn:
            result = conn.execute(
                sql_text("SELECT storage_path, filename, file_type FROM document_kb_documents WHERE id = :id"),
                {"id": int(file_id) if file_id.isdigit() else 0},
            ).fetchone()
            if result and result[0]:
                stored_file = Path("/app/data/documents") / result[0]
                if stored_file.exists():
                    return _serve_file_with_viewer_conversion(stored_file, original_filename=result[1])

        raise HTTPException(status_code=404, detail="File not found")

    # --- Delete Document ---
    @app.delete("/api/v1/documents/{doc_id}", tags=["Documents"])
    async def delete_kb_document(doc_id: str, kb_id: str = "", workspace_id: str = ""):
        """Delete a document and its chunks from the knowledge base."""
        from sqlalchemy import text as sql_text

        engine = _get_doc_db_engine()
        with engine.connect() as conn:
            conn.execute(sql_text("DELETE FROM document_kb_chunks WHERE document_id = :id"), {"id": int(doc_id)})
            conn.execute(sql_text("DELETE FROM document_kb_documents WHERE id = :id"), {"id": int(doc_id)})
            conn.commit()

        return {"status": "deleted", "doc_id": doc_id}

    # ==========================================================================
    # Metrics
    # ==========================================================================

    @app.get("/metrics", tags=["Monitoring"])
    async def prometheus_metrics() -> str:
        """Prometheus metrics endpoint."""
        from prometheus_client import generate_latest, CONTENT_TYPE_LATEST
        from starlette.responses import Response

        return Response(
            content=generate_latest(),
            media_type=CONTENT_TYPE_LATEST,
        )
